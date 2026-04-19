"""
Generador de documentación ejecutiva — SFTP Secure Ticket Service
Genera un .docx profesional de 50-100 páginas siguiendo el índice del TFM.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── Estilos base ───────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x3A, 0x28)
    hs.font.name = 'Calibri'

doc.styles['Heading 1'].font.size = Pt(18)
doc.styles['Heading 2'].font.size = Pt(14)
doc.styles['Heading 3'].font.size = Pt(12)


def add_code_block(text, doc=doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    # shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    run._element.rPr.append(shading)
    return p


def add_table(headers, rows, doc=doc):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()
    return table


def page_break():
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SFTP SECURE TICKET SERVICE')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Sistema de Transferencia Segura de Ficheros\ncon Autenticación MFA y Tickets de Un Solo Uso')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Ejecutiva del Proyecto')
run.font.size = Pt(13)
run.font.italic = True

doc.add_paragraph()

info_items = [
    ('Versión', '1.0'),
    ('Fecha', 'Abril 2026'),
    ('Clasificación', 'Confidencial — Uso Interno'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(value)
    run.font.size = Pt(11)

page_break()

# ═══════════════════════════════════════════════════════════════
#  AGRADECIMIENTOS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Agradecimientos', level=1)
doc.add_paragraph()

p = doc.add_paragraph(
    'A mi tutor/a de Trabajo de Fin de Máster, por su orientación, paciencia y rigor '
    'académico a lo largo de todo el proceso de desarrollo de este proyecto.'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    'A los profesores del Máster, cuyas enseñanzas en criptografía, seguridad de sistemas '
    'y arquitecturas de red han sido el fundamento teórico sobre el que se construye este trabajo.'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    'A mi familia y amigos, por su apoyo incondicional durante estos meses de trabajo intenso, '
    'por comprender las ausencias y celebrar los avances.'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    'A la comunidad open-source, por mantener y evolucionar las herramientas que hacen posible '
    'proyectos como este: Python, Flask, Docker, OpenLDAP, Redis, Nginx y las decenas de '
    'bibliotecas que conforman el ecosistema.'
)
p.paragraph_format.space_after = Pt(12)

p = doc.add_paragraph(
    'Y a todos aquellos que, desde distintos ámbitos, contribuyen a hacer de Internet '
    'un lugar más seguro.'
)
p.italic = True

page_break()

# ═══════════════════════════════════════════════════════════════
#  RESUMEN / ABSTRACT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Resumen', level=1)

doc.add_paragraph(
    'Este Trabajo de Fin de Máster presenta el diseño, implementación y despliegue de un sistema '
    'completo de transferencia segura de ficheros denominado SFTP Secure Ticket Service. El sistema '
    'integra, de forma nativa y sin dependencias de terceros, múltiples capas de seguridad: '
    'autenticación multifactor (MFA) combinando credenciales LDAP con códigos TOTP, autorización '
    'mediante tickets de un solo uso (OTP) con caducidad temporal, firma digital RSA-4096 con '
    'esquema PSS-SHA512 de cada fichero subido, verificación obligatoria de integridad en cada '
    'descarga y detección automática de ficheros comprometidos con apertura de incidentes de seguridad.'
)

doc.add_paragraph(
    'La arquitectura se basa en microservicios contenedorizados con Docker Compose: un proxy inverso '
    'Nginx con terminación TLS, un backend Flask (Python 3.12) que expone una API REST, un directorio '
    'OpenLDAP para la gestión de identidades, Redis como almacén de tickets y metadatos, y MailHog '
    'como servidor SMTP de desarrollo. El frontend es una SPA vanilla (HTML/CSS/JS) con estética de '
    'terminal de seguridad que no depende de frameworks ni de node_modules, eliminando riesgos de '
    'supply chain attack.'
)

doc.add_paragraph(
    'El sistema sigue la filosofía Zero Trust: cada operación de subida o descarga requiere un ticket '
    'OTP individual, válido durante 5 minutos y de un solo uso. Cada fichero subido se firma '
    'digitalmente y se genera un recibo digital verificable que se envía al usuario por correo '
    'electrónico. En la descarga, se verifica criptográficamente que el fichero no ha sido alterado; '
    'si se detecta una discrepancia, la descarga se bloquea y se abre un incidente automáticamente.'
)

doc.add_paragraph(
    'El diseño cumple con los requisitos del Esquema Nacional de Seguridad (ENS), el RGPD, '
    'ISO/IEC 27001:2022 y las recomendaciones del NIST (SP 800-57, SP 800-61, SP 800-63B).'
)

p = doc.add_paragraph()
run = p.add_run('Palabras clave: ')
run.bold = True
p.add_run('transferencia segura de ficheros, autenticación multifactor, TOTP, tickets OTP, '
          'firma digital RSA-4096, PSS-SHA512, verificación de integridad, Zero Trust, Docker, '
          'Flask, OpenLDAP, Redis, gestión de incidentes.')

doc.add_paragraph()
doc.add_paragraph()

# ── Abstract (English) ──
doc.add_heading('Abstract', level=1)

doc.add_paragraph(
    'This Master\'s Thesis presents the design, implementation and deployment of a complete secure '
    'file transfer system called SFTP Secure Ticket Service. The system natively integrates multiple '
    'security layers without third-party dependencies: multi-factor authentication (MFA) combining '
    'LDAP credentials with TOTP codes, authorization through one-time-use tickets (OTP) with '
    'configurable expiry, RSA-4096 digital signatures using the PSS-SHA512 scheme for every uploaded '
    'file, mandatory integrity verification on every download, and automatic detection of compromised '
    'files with automated incident creation.'
)

doc.add_paragraph(
    'The architecture is based on containerized microservices orchestrated with Docker Compose: '
    'an Nginx reverse proxy with TLS termination, a Flask backend (Python 3.12) exposing a REST API, '
    'an OpenLDAP directory for identity management, Redis as a ticket and metadata store, and MailHog '
    'as a development SMTP server. The frontend is a vanilla SPA (HTML/CSS/JS) with a security '
    'terminal aesthetic that relies on no frameworks or node_modules, eliminating supply chain '
    'attack risks.'
)

doc.add_paragraph(
    'The system follows a Zero Trust philosophy: every upload or download operation requires an '
    'individual OTP ticket, valid for 5 minutes and single-use only. Every uploaded file is '
    'digitally signed and a verifiable digital receipt is generated and sent to the user via email. '
    'On download, the file\'s integrity is cryptographically verified; if a discrepancy is detected, '
    'the download is blocked and an incident is automatically opened.'
)

doc.add_paragraph(
    'The design complies with the Spanish National Security Framework (ENS), GDPR, '
    'ISO/IEC 27001:2022 and NIST recommendations (SP 800-57, SP 800-61, SP 800-63B).'
)

p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.bold = True
p.add_run('secure file transfer, multi-factor authentication, TOTP, OTP tickets, '
          'RSA-4096 digital signatures, PSS-SHA512, integrity verification, Zero Trust, Docker, '
          'Flask, OpenLDAP, Redis, incident management.')

page_break()

# ═══════════════════════════════════════════════════════════════
#  ÍNDICE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Índice de Contenidos', level=1)
doc.add_paragraph()

toc_items = [
    'Agradecimientos',
    'Resumen / Abstract',
    '',
    '1. Introducción',
    '   1.1 Contexto',
    '   1.2 Objetivos del proyecto',
    '   1.3 Alcance y limitaciones',
    '2. Estado del arte',
    '   2.1 Protocolos de transferencia segura de ficheros',
    '   2.2 Autenticación multifactor (MFA) en entornos corporativos',
    '   2.3 Sistemas de tickets y tokens de un solo uso (OTP)',
    '   2.4 Firma digital y verificación de integridad',
    '   2.5 Estándares y normativa aplicable',
    '3. Análisis de requisitos',
    '   3.1 Requisitos funcionales',
    '   3.2 Requisitos de seguridad',
    '   3.3 Requisitos de despliegue',
    '   3.4 Casos de uso',
    '4. Diseño de la arquitectura',
    '   4.1 Arquitectura general del sistema',
    '   4.2 Stack tecnológico y justificación',
    '   4.3 Modelo de autenticación (LDAP + TOTP)',
    '   4.4 Modelo de autorización basado en tickets OTP',
    '   4.5 Modelo de firma digital y registro de integridad',
    '   4.6 Diagrama de flujo de operaciones',
    '   4.7 Modelo de datos (Redis, LDAP)',
    '5. Implementación',
    '   5.1 Contenedorización con Docker',
    '   5.2 Backend Flask: API REST',
    '   5.3 Servicio de autenticación LDAP',
    '   5.4 Generación de QR y verificación TOTP',
    '   5.5 Gestión de tickets OTP en Redis',
    '   5.6 Firma digital RSA-4096 (PSS-SHA512)',
    '   5.7 Verificación de integridad en descarga',
    '   5.8 Notificaciones por correo electrónico',
    '   5.9 Frontend: interfaz terminal web',
    '   5.10 Proxy inverso y terminación TLS (Nginx)',
    '6. Despliegue seguro',
    '   6.1 Entorno de despliegue (Ubuntu Server)',
    '   6.2 Configuración de red y firewalls',
    '   6.3 Gestión de secretos y variables de entorno',
    '   6.4 Certificados TLS',
    '   6.5 Bastionado de contenedores',
    '   6.6 Logs y auditoría',
    '   6.7 Plan de mantenimiento y actualizaciones',
    '7. Gestión de incidentes',
    '   7.1 Modelo de detección de ficheros comprometidos',
    '   7.2 Protocolo de respuesta a incidentes (NIST SP 800-61)',
    '   7.3 Fases: contención, análisis, erradicación, recuperación y post-mortem',
    '   7.4 Diferenciación de vectores de ataque',
    '   7.5 Simulación de incidente real',
    '8. Pruebas',
    '   8.1 Pruebas de autenticación MFA',
    '   8.2 Pruebas de subida y descarga con tickets',
    '   8.3 Pruebas de verificación de integridad',
    '   8.4 Pruebas de detección de ficheros comprometidos',
    '   8.5 Pruebas de rendimiento y límites',
    '9. Conclusiones y trabajo futuro',
    '   9.1 Objetivos alcanzados',
    '   9.2 Limitaciones conocidas',
    '   9.3 Mejoras futuras',
    '   9.4 Valoración personal',
    '',
    'Anexo A: Glosario de términos',
    'Anexo B: Referencias bibliográficas',
    'Anexo C: Capturas de pantalla del sistema',
    'Anexo D: Código fuente',
]
for item in toc_items:
    if item == '':
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if not item.startswith('   '):
        p.runs[0].bold = True

page_break()

# ═══════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ═══════════════════════════════════════════════════════════════

doc.add_heading('1. Introducción', level=1)

doc.add_heading('1.1 Contexto', level=2)

doc.add_paragraph(
    'La transformación digital de las organizaciones ha traído consigo un incremento exponencial '
    'en el volumen de datos que se intercambian entre sistemas, departamentos y entidades externas. '
    'Los ficheros que transitan por estas infraestructuras contienen frecuentemente información '
    'clasificada, datos personales protegidos por el RGPD, documentación financiera sujeta a '
    'auditoría o propiedad intelectual de alto valor estratégico.'
)

doc.add_paragraph(
    'En este contexto, las soluciones tradicionales de transferencia de ficheros —basadas en '
    'FTP sin cifrar, carpetas compartidas por SMB o incluso el envío de adjuntos por correo '
    'electrónico— presentan vulnerabilidades críticas: ausencia de cifrado en tránsito, falta '
    'de trazabilidad sobre quién accede a qué fichero, imposibilidad de verificar que un '
    'documento no ha sido alterado tras su almacenamiento y carencia de mecanismos de '
    'autenticación robustos que vayan más allá del par usuario/contraseña.'
)

doc.add_paragraph(
    'Los incidentes de seguridad relacionados con la transferencia de ficheros no son anecdóticos. '
    'Según el informe de Verizon Data Breach Investigations Report (DBIR) de 2024, el 15% de las '
    'brechas de datos en entornos corporativos involucraron la exfiltración o manipulación de '
    'ficheros en tránsito o en reposo. El coste medio de una brecha de datos alcanzó los 4,45 '
    'millones de dólares según IBM Security en 2023, con un tiempo medio de detección de 204 días. '
    'Estos datos subrayan la necesidad de implementar controles de seguridad en capas (defense in '
    'depth) que no se limiten al cifrado del canal de comunicación.'
)

doc.add_paragraph(
    'El presente proyecto nace de la necesidad real de dotar a una organización de un sistema '
    'de transferencia segura de ficheros que integre, de forma nativa y sin dependencias de '
    'terceros, las siguientes capacidades de seguridad: autenticación multifactor (MFA) con '
    'verificación TOTP, autorización mediante tickets de un solo uso (OTP) con caducidad '
    'temporal, firma digital RSA-4096 de cada fichero subido, verificación obligatoria de '
    'integridad en cada descarga y detección automática de ficheros comprometidos con apertura '
    'de incidentes de seguridad.'
)

doc.add_paragraph(
    'El sistema se ha diseñado siguiendo los principios de Zero Trust Architecture, donde '
    'ninguna operación se considera segura por defecto: cada acción (subida, descarga, '
    'verificación) requiere una autorización explícita, temporal y verificable. Esta filosofía '
    'se materializa en un flujo de operaciones que exige, para cada transferencia, la generación '
    'de un ticket OTP válido durante 5 minutos y de un solo uso, garantizando que incluso si '
    'un token de sesión fuese comprometido, el atacante no podría realizar operaciones sin '
    'obtener adicionalmente un ticket válido.'
)

doc.add_heading('1.2 Objetivos del proyecto', level=2)

doc.add_paragraph(
    'El objetivo principal de este proyecto es diseñar, implementar y documentar un sistema '
    'completo de transferencia segura de ficheros que cumpla con los más altos estándares de '
    'seguridad de la información, manteniendo al mismo tiempo una experiencia de usuario '
    'fluida y una arquitectura fácilmente desplegable mediante contenedores Docker.'
)

doc.add_paragraph('Los objetivos específicos del proyecto son:')

objectives = [
    ('OBJ-01', 'Implementar autenticación en dos factores (2FA) combinando credenciales LDAP con códigos TOTP (Time-based One-Time Password) compatibles con Google Authenticator y Authy.'),
    ('OBJ-02', 'Diseñar un sistema de tickets OTP de un solo uso con caducidad configurable para autorizar individualmente cada operación de subida o descarga de ficheros.'),
    ('OBJ-03', 'Integrar firma digital RSA-4096 con esquema PSS-SHA512 para cada fichero subido, generando un recibo digital verificable que sirva como evidencia de integridad.'),
    ('OBJ-04', 'Implementar verificación obligatoria de integridad en cada descarga, comparando el hash SHA-512 actual del fichero con el hash firmado en el momento de la subida.'),
    ('OBJ-05', 'Desarrollar un mecanismo automático de detección de ficheros comprometidos que, ante una discrepancia de hash o firma, bloquee la descarga y abra un incidente de seguridad.'),
    ('OBJ-06', 'Generar un registro de auditoría completo (audit log) de todas las operaciones del sistema: autenticaciones, generación de tickets, subidas, descargas, verificaciones e incidentes.'),
    ('OBJ-07', 'Enviar notificaciones por correo electrónico con el recibo digital tras cada subida de fichero, proporcionando al usuario el hash necesario para futuras descargas.'),
    ('OBJ-08', 'Desplegar el sistema completo mediante Docker Compose con configuración de seguridad de producción: TLS, rate limiting, headers de seguridad, contenedores sin privilegios.'),
    ('OBJ-09', 'Documentar el sistema de forma exhaustiva, incluyendo análisis de requisitos, diseño arquitectónico, decisiones de implementación, guía de despliegue y plan de respuesta a incidentes.'),
]

add_table(
    ['ID', 'Descripción del objetivo'],
    [[o[0], o[1]] for o in objectives]
)

doc.add_heading('1.3 Alcance y limitaciones', level=2)

doc.add_paragraph('Alcance del proyecto:')

scope_items = [
    'Diseño e implementación completa de una API REST segura con Flask para la gestión del ciclo de vida de ficheros.',
    'Integración con servidor LDAP (OpenLDAP) para la gestión centralizada de identidades y credenciales.',
    'Implementación de autenticación MFA con generación de códigos QR para enrollment TOTP.',
    'Sistema de tickets OTP almacenados en Redis con TTL configurable y semántica de un solo uso.',
    'Firma digital RSA-4096 con esquema PSS y hash SHA-512, incluyendo generación automática de claves.',
    'Frontend SPA (Single Page Application) con interfaz estilo terminal que cubre todo el flujo funcional.',
    'Infraestructura de despliegue con Docker Compose: Nginx (reverse proxy + TLS), Flask (backend), OpenLDAP, Redis y MailHog.',
    'Sistema de notificaciones por email con recibos digitales en formato HTML y texto plano.',
    'API de gestión de incidentes de seguridad con detección automática de ficheros comprometidos.',
    'Documentación ejecutiva completa del proyecto.',
]
for item in scope_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Limitaciones del proyecto:')

limitations = [
    'Los ficheros se almacenan en el sistema de archivos del servidor sin cifrado at-rest. El cifrado AES-256 de ficheros en reposo se contempla como mejora futura.',
    'La clave RSA de firma se almacena en un fichero PEM en el servidor. En un entorno de producción crítico, debería utilizarse un HSM (Hardware Security Module).',
    'El servicio SMTP de desarrollo utiliza MailHog (sin entrega real). En producción debe configurarse un relay SMTP autenticado.',
    'No se implementa clustering ni alta disponibilidad. El sistema está diseñado para despliegue en un solo nodo Docker.',
    'El directorio LDAP se inicializa con usuarios de prueba. En producción se integraría con el directorio corporativo existente (Active Directory, FreeIPA, etc.).',
    'No se implementa rotación automática de la clave RSA. Se recomienda rotación manual trimestral con re-firma de ficheros existentes.',
    'La interfaz web no soporta internacionalización (i18n). Todo el contenido está en español.',
]
for item in limitations:
    doc.add_paragraph(item, style='List Bullet')

page_break()

# ═══════════════════════════════════════════════════════════════
#  2. ESTADO DEL ARTE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('2. Estado del arte', level=1)

doc.add_paragraph(
    'En esta sección se revisan las tecnologías, protocolos y estándares que constituyen la '
    'base teórica y práctica del sistema desarrollado. El objetivo es situar el proyecto en el '
    'contexto actual de la seguridad de la información y justificar las decisiones de diseño '
    'adoptadas.'
)

doc.add_heading('2.1 Protocolos de transferencia segura de ficheros (SFTP, FTPS, SCP)', level=2)

doc.add_paragraph(
    'La transferencia de ficheros entre sistemas ha evolucionado significativamente desde los '
    'primeros protocolos FTP (File Transfer Protocol), definido en el RFC 959 en 1985. FTP '
    'transmite tanto credenciales como datos en texto claro, lo que lo hace completamente '
    'inadecuado para entornos donde la confidencialidad es un requisito.'
)

doc.add_paragraph('Los principales protocolos de transferencia segura utilizados actualmente son:')

doc.add_heading('SFTP (SSH File Transfer Protocol)', level=3)

doc.add_paragraph(
    'SFTP opera como un subsistema del protocolo SSH (Secure Shell) y proporciona transferencia '
    'de ficheros cifrada sobre un canal SSH. A diferencia de FTP, SFTP utiliza un único puerto '
    '(por defecto el 22) y cifra tanto la autenticación como la transferencia de datos. Soporta '
    'operaciones completas de gestión de ficheros: lectura, escritura, eliminación, cambio de '
    'permisos y listado de directorios.'
)

doc.add_paragraph(
    'Ventajas: cifrado completo del canal, autenticación por clave pública, uso de un solo '
    'puerto (simplifica la configuración de firewalls), amplio soporte en clientes y servidores. '
    'Limitaciones: no proporciona verificación de integridad a nivel de fichero individual '
    '(solo integridad del canal SSH), no incluye mecanismos nativos de autorización granular '
    'por operación ni trazabilidad de acceso.'
)

doc.add_heading('FTPS (FTP Secure / FTP sobre TLS)', level=3)

doc.add_paragraph(
    'FTPS extiende FTP añadiendo soporte para cifrado TLS/SSL. Existen dos variantes: FTPS '
    'explícito (el cliente solicita el cifrado tras conectar al puerto 21) y FTPS implícito '
    '(el cifrado se inicia desde la conexión en el puerto 990). FTPS mantiene la separación '
    'entre canal de control y canal de datos propia de FTP, lo que requiere la apertura de '
    'múltiples puertos en el firewall y complica el NAT traversal.'
)

doc.add_paragraph(
    'Ventajas: cifrado TLS estándar, soporte de certificados X.509, compatibilidad con '
    'infraestructuras PKI existentes. Limitaciones: complejidad de configuración de firewall, '
    'uso de múltiples puertos, menor adopción que SFTP en entornos Linux/Unix.'
)

doc.add_heading('SCP (Secure Copy Protocol)', level=3)

doc.add_paragraph(
    'SCP utiliza SSH para la transferencia segura de ficheros, pero con funcionalidad más '
    'limitada que SFTP: solo permite copiar ficheros, sin capacidad de listar directorios, '
    'eliminar ficheros o cambiar permisos. Es considerado un protocolo legacy y se recomienda '
    'usar SFTP en su lugar.'
)

doc.add_paragraph(
    'El sistema desarrollado en este proyecto no implementa el protocolo SFTP a nivel de '
    'transporte, sino que utiliza HTTPS con TLS 1.2/1.3 como capa de transporte cifrada y '
    'construye sobre ella un sistema de transferencia basado en API REST con mecanismos de '
    'seguridad adicionales (MFA, tickets OTP, firma digital) que los protocolos tradicionales '
    'no ofrecen de forma nativa. El nombre "SFTP Secure Ticket Service" hace referencia al '
    'concepto de transferencia segura de ficheros, no a la implementación del protocolo SFTP '
    'en sentido estricto.'
)

doc.add_heading('Comparativa de protocolos', level=3)

add_table(
    ['Característica', 'FTP', 'FTPS', 'SCP', 'SFTP', 'Este sistema'],
    [
        ['Cifrado en tránsito', 'No', 'TLS', 'SSH', 'SSH', 'TLS 1.2/1.3'],
        ['Autenticación MFA', 'No', 'No', 'No', 'Posible*', 'Sí (LDAP+TOTP)'],
        ['Autorización por operación', 'No', 'No', 'No', 'No', 'Sí (Tickets OTP)'],
        ['Firma digital de ficheros', 'No', 'No', 'No', 'No', 'Sí (RSA-4096)'],
        ['Verificación integridad', 'No', 'No', 'No', 'Canal', 'Por fichero (SHA-512)'],
        ['Detección de compromiso', 'No', 'No', 'No', 'No', 'Sí (automática)'],
        ['Audit log nativo', 'No', 'No', 'No', 'Parcial', 'Sí (completo)'],
        ['Notificación por email', 'No', 'No', 'No', 'No', 'Sí (con recibo)'],
    ]
)

doc.add_paragraph('* SFTP puede configurarse con MFA a nivel de SSH, pero no es una funcionalidad nativa del protocolo de transferencia.')

doc.add_heading('2.2 Autenticación multifactor (MFA) en entornos corporativos', level=2)

doc.add_paragraph(
    'La autenticación multifactor (MFA) es un mecanismo de seguridad que requiere que el '
    'usuario presente al menos dos tipos de evidencia (factores) de su identidad antes de '
    'obtener acceso. Los factores se clasifican en tres categorías clásicas:'
)

factors = [
    ('Algo que sabes (knowledge)', 'Contraseña, PIN, respuesta a pregunta de seguridad.'),
    ('Algo que tienes (possession)', 'Token hardware (YubiKey), teléfono móvil con app de autenticación, tarjeta inteligente.'),
    ('Algo que eres (inherence)', 'Huella dactilar, reconocimiento facial, escaneo de iris.'),
]
add_table(['Factor', 'Ejemplos'], [[f[0], f[1]] for f in factors])

doc.add_paragraph(
    'El sistema implementado utiliza una combinación de dos factores: conocimiento '
    '(contraseña LDAP) y posesión (código TOTP generado por una aplicación móvil). Esta '
    'combinación ofrece un equilibrio óptimo entre seguridad y usabilidad:'
)

doc.add_paragraph(
    'TOTP (Time-based One-Time Password) es un algoritmo definido en el RFC 6238 que genera '
    'códigos numéricos de 6 dígitos a partir de un secreto compartido y el tiempo actual. El '
    'código cambia cada 30 segundos y es válido durante una ventana configurable (en nuestro '
    'sistema, un período de 30 segundos ± 1 período adicional para compensar desincronización '
    'de reloj). El algoritmo utiliza HMAC-SHA1 internamente, aunque variantes más recientes '
    'soportan SHA-256 y SHA-512.'
)

doc.add_paragraph(
    'El flujo de enrollment MFA en el sistema es el siguiente: tras la primera autenticación '
    'LDAP exitosa, el servidor genera un secreto TOTP aleatorio de 160 bits (base32), lo '
    'almacena en el campo description del usuario LDAP y genera un código QR con el URI de '
    'provisioning (otpauth://totp/...). El usuario escanea el QR con su aplicación de '
    'autenticación (Google Authenticator, Authy, Microsoft Authenticator) y a partir de ese '
    'momento puede generar códigos TOTP válidos para completar el segundo factor.'
)

doc.add_paragraph(
    'En comparación con otros mecanismos de segundo factor (SMS, email, push notifications, '
    'WebAuthn/FIDO2), TOTP ofrece las siguientes ventajas para este caso de uso: no requiere '
    'conexión a internet en el dispositivo del usuario, no depende de servicios de terceros '
    '(como una pasarela SMS), es compatible con múltiples aplicaciones de autenticación y su '
    'implementación es relativamente simple. Como limitación, TOTP es vulnerable a ataques '
    'de phishing en tiempo real (el atacante puede interceptar el código y usarlo '
    'inmediatamente), aunque esto se mitiga con la corta ventana de validez del código.'
)

doc.add_heading('2.3 Sistemas de tickets y tokens de un solo uso (OTP)', level=2)

doc.add_paragraph(
    'Los sistemas de tickets de un solo uso (One-Time Password / One-Time Token) son '
    'mecanismos de autorización que generan credenciales temporales válidas para una única '
    'operación. Una vez utilizado, el ticket queda invalidado de forma permanente, '
    'independientemente de si la operación tuvo éxito o no.'
)

doc.add_paragraph(
    'Este concepto tiene su origen en los sistemas Kerberos, donde los "tickets" son tokens '
    'criptográficos emitidos por un Key Distribution Center (KDC) que autorizan el acceso a '
    'servicios específicos durante un tiempo limitado. El sistema implementado en este proyecto '
    'adapta este concepto a un contexto más simple pero igualmente robusto:'
)

doc.add_paragraph(
    'Un ticket OTP en nuestro sistema es un token criptográficamente aleatorio de 48 bytes '
    '(generado con secrets.token_urlsafe) que se almacena en Redis con un TTL (Time To Live) '
    'de 300 segundos (5 minutos) por defecto. El ticket contiene metadatos sobre la operación '
    'autorizada (upload/download), el usuario que lo solicitó, la IP de origen y el nombre '
    'del fichero (en caso de descarga). Cuando el ticket se consume, se marca como usado y '
    'se mantiene brevemente en Redis (60 segundos) para la auditoría antes de ser eliminado '
    'automáticamente.'
)

doc.add_paragraph(
    'Las ventajas de este modelo de autorización son:'
)

otp_advantages = [
    'Granularidad: cada operación requiere su propio ticket, lo que impide que un token de sesión comprometido sea suficiente para realizar operaciones.',
    'Caducidad temporal: los tickets expiran automáticamente, reduciendo la ventana de explotación en caso de intercepción.',
    'Un solo uso: incluso si un atacante captura un ticket en tránsito, no puede reutilizarlo si ya fue consumido.',
    'Vinculación a operación: un ticket de subida no puede usarse para descargar, y viceversa.',
    'Audit trail: la creación, consumo e invalidación de cada ticket queda registrada en el log de auditoría.',
    'Desacoplamiento: el sistema de tickets es independiente del mecanismo de autenticación, lo que permite cambiar la capa de autenticación sin modificar la lógica de autorización.',
]
for item in otp_advantages:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.4 Firma digital y verificación de integridad', level=2)

doc.add_paragraph(
    'La firma digital es un mecanismo criptográfico que permite verificar la autenticidad '
    'e integridad de un documento digital. A diferencia del cifrado (que protege la '
    'confidencialidad), la firma digital garantiza dos propiedades:'
)

doc.add_paragraph(
    'Integridad: el documento no ha sido modificado desde que fue firmado. '
    'Autenticidad: el documento fue firmado por la entidad que posee la clave privada.'
)

doc.add_paragraph(
    'El sistema utiliza RSA-4096 con esquema de firma PSS (Probabilistic Signature Scheme) '
    'y función hash SHA-512. Esta combinación se considera criptográficamente robusta según '
    'las recomendaciones actuales del NIST (SP 800-57 Part 1, Rev. 5) y la ANSSI francesa. '
    'El flujo de firma e integridad del sistema es:'
)

doc.add_paragraph(
    '1. Subida del fichero: el sistema calcula el hash SHA-512 del fichero completo, '
    'firma el hash con la clave privada RSA-4096 usando PSS-SHA512, genera un recibo '
    'digital con el hash y la firma, almacena los metadatos en Redis y envía el recibo '
    'al usuario por email.'
)

doc.add_paragraph(
    '2. Descarga del fichero: el usuario debe proporcionar el hash SHA-512 original '
    '(obtenido del recibo de subida). El sistema recalcula el hash del fichero en disco '
    'y lo compara con el hash proporcionado. Si no coinciden, la descarga se bloquea y '
    'se abre un incidente de seguridad automáticamente. Adicionalmente, se verifica la '
    'firma RSA almacenada contra el hash actual del fichero.'
)

doc.add_paragraph(
    'El esquema PSS (Probabilistic Signature Scheme) se elige sobre PKCS#1 v1.5 porque '
    'PSS tiene una demostración de seguridad formal (es demostrado seguro en el modelo '
    'del oráculo aleatorio), mientras que PKCS#1 v1.5 tiene vulnerabilidades conocidas '
    '(ataque de Bleichenbacher). PSS introduce un componente aleatorio (salt) en cada '
    'operación de firma, lo que significa que firmar el mismo mensaje dos veces produce '
    'firmas diferentes, dificultando ciertos ataques de análisis.'
)

doc.add_heading('RSA-4096 frente a alternativas', level=3)

add_table(
    ['Algoritmo', 'Tamaño clave', 'Seguridad equiv. (bits)', 'Rendimiento firma', 'Rendimiento verificación', 'Post-cuántico'],
    [
        ['RSA-2048', '2048 bits', '~112', 'Medio', 'Rápido', 'No'],
        ['RSA-4096', '4096 bits', '~140', 'Lento', 'Rápido', 'No'],
        ['ECDSA P-256', '256 bits', '~128', 'Rápido', 'Medio', 'No'],
        ['ECDSA P-384', '384 bits', '~192', 'Rápido', 'Medio', 'No'],
        ['Ed25519', '256 bits', '~128', 'Muy rápido', 'Muy rápido', 'No'],
        ['CRYSTALS-Dilithium', '~2.5 KB', '~128 (PQ)', 'Rápido', 'Rápido', 'Sí'],
    ]
)

doc.add_paragraph(
    'Se elige RSA-4096 por su amplia compatibilidad con bibliotecas criptográficas estándar, '
    'su madurez (más de 40 años de criptoanálisis) y el requisito del proyecto de usar un '
    'tamaño de clave que garantice seguridad a largo plazo. ECDSA (Elliptic Curve Digital '
    'Signature Algorithm) sería una alternativa válida con mejor rendimiento, pero RSA-4096 '
    'ofrece un margen de seguridad mayor y es más ampliamente entendido en auditorías de '
    'cumplimiento normativo. La migración a ECDSA o algoritmos post-cuánticos se contempla '
    'como trabajo futuro.'
)

doc.add_heading('2.5 Estándares y normativa aplicable (ENS, RGPD, ISO 27001, NIST)', level=2)

doc.add_paragraph(
    'El diseño del sistema tiene en cuenta los principales marcos normativos y estándares '
    'de seguridad aplicables en el contexto europeo y español:'
)

doc.add_heading('Esquema Nacional de Seguridad (ENS) — RD 311/2022', level=3)

doc.add_paragraph(
    'El ENS establece los principios básicos y requisitos mínimos para la protección de la '
    'información en el sector público español. El sistema contribuye al cumplimiento de las '
    'siguientes medidas del ENS: mp.info.2 (calificación de la información), mp.info.3 '
    '(firma electrónica), mp.com.2 (protección de la confidencialidad en tránsito), '
    'mp.s.2 (protección de servicios), op.acc.5 (mecanismo de autenticación — MFA), '
    'op.acc.6 (acceso local — control de sesiones con JWT), op.exp.8 (registro de '
    'actividad — audit log), op.exp.10 (protección de claves criptográficas).'
)

doc.add_heading('Reglamento General de Protección de Datos (RGPD)', level=3)

doc.add_paragraph(
    'El RGPD (Reglamento UE 2016/679) exige la implementación de medidas técnicas y '
    'organizativas apropiadas para garantizar la seguridad del tratamiento de datos '
    'personales. El sistema contribuye al cumplimiento del artículo 32 (seguridad del '
    'tratamiento) mediante: cifrado de datos en tránsito (TLS 1.2/1.3), control de '
    'acceso basado en autenticación MFA, registro de todas las operaciones de acceso '
    'a datos, capacidad de verificación de integridad de los datos almacenados y '
    'notificación automática al usuario de cualquier operación sobre sus ficheros.'
)

doc.add_heading('ISO/IEC 27001:2022', level=3)

doc.add_paragraph(
    'ISO 27001 define los requisitos para un Sistema de Gestión de Seguridad de la '
    'Información (SGSI). Los controles del Anexo A implementados por el sistema incluyen: '
    'A.8.2 (Gestión de acceso privilegiado), A.8.3 (Restricción de acceso a la información), '
    'A.8.5 (Autenticación segura — MFA), A.8.9 (Gestión de la configuración), A.8.15 '
    '(Registro de eventos — audit log), A.8.20 (Seguridad de redes), A.8.24 (Uso de '
    'criptografía) y A.8.25 (Ciclo de vida del desarrollo seguro).'
)

doc.add_heading('NIST Cybersecurity Framework y publicaciones especiales', level=3)

doc.add_paragraph(
    'El NIST (National Institute of Standards and Technology) de EE.UU. proporciona guías '
    'técnicas de referencia mundial. El sistema sigue las recomendaciones de: SP 800-63B '
    '(Digital Identity Guidelines — Authentication), SP 800-57 (Recommendation for Key '
    'Management — elección de RSA-4096), SP 800-61 Rev. 2 (Computer Security Incident '
    'Handling Guide — protocolo de gestión de incidentes), y el NIST Cybersecurity Framework '
    '2.0 (funciones Identify, Protect, Detect, Respond, Recover).'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  3. ANÁLISIS DE REQUISITOS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('3. Análisis de requisitos', level=1)

doc.add_heading('3.1 Requisitos funcionales', level=2)

doc.add_paragraph(
    'A continuación se detallan los requisitos funcionales del sistema, agrupados por '
    'áreas funcionales.'
)

rf_data = [
    ['RF-01', 'Autenticación LDAP', 'El sistema debe autenticar usuarios contra un directorio LDAP mediante binding con credenciales (uid + contraseña).'],
    ['RF-02', 'Enrollment MFA', 'Tras la primera autenticación exitosa, el sistema debe generar un secreto TOTP, codificarlo en un QR y presentarlo al usuario para su escaneo con una app de autenticación.'],
    ['RF-03', 'Verificación TOTP', 'El sistema debe verificar códigos TOTP de 6 dígitos con una ventana de tolerancia de ±1 período (30 segundos).'],
    ['RF-04', 'Emisión de JWT', 'Tras completar la autenticación MFA, el sistema debe emitir un token JWT firmado con HS256, con caducidad configurable (por defecto 15 minutos).'],
    ['RF-05', 'Blacklist de tokens', 'Al cerrar sesión, el token JWT debe añadirse a una blacklist en Redis para impedir su reutilización.'],
    ['RF-06', 'Generación de tickets OTP', 'El usuario autenticado debe poder solicitar tickets de un solo uso para operaciones de subida o descarga, con TTL configurable.'],
    ['RF-07', 'Subida de ficheros', 'El sistema debe aceptar ficheros de hasta 500 MB, almacenarlos en el directorio del usuario y generar hash SHA-512 + firma RSA.'],
    ['RF-08', 'Recibo digital', 'Tras cada subida, el sistema debe generar un recibo digital estructurado con ID, timestamp, hash, firma, algoritmo y estado de integridad.'],
    ['RF-09', 'Notificación email', 'El sistema debe enviar un correo electrónico al usuario con el recibo digital en formato HTML y texto plano.'],
    ['RF-10', 'Descarga con verificación', 'Para descargar un fichero, el usuario debe proporcionar el hash SHA-512 original. El sistema recalcula el hash y solo permite la descarga si coincide.'],
    ['RF-11', 'Detección de compromiso', 'Si el hash del fichero no coincide con el almacenado o proporcionado, el sistema debe bloquear la descarga y registrar un incidente de seguridad.'],
    ['RF-12', 'Listado de ficheros', 'El usuario debe poder listar sus ficheros con metadatos: nombre, tamaño, fecha de subida, hash parcial e ID de recibo.'],
    ['RF-13', 'Verificación de recibo', 'El usuario debe poder verificar manualmente la integridad de un fichero proporcionando su hash SHA-512.'],
    ['RF-14', 'Gestión de incidentes', 'El sistema debe mantener un registro de incidentes de seguridad consultable vía API.'],
    ['RF-15', 'Health check', 'El sistema debe exponer un endpoint de salud que verifique la conectividad con Redis y LDAP.'],
]

add_table(['ID', 'Nombre', 'Descripción'], rf_data)

doc.add_heading('3.2 Requisitos de seguridad', level=2)

rs_data = [
    ['RS-01', 'TLS obligatorio', 'Toda comunicación debe realizarse sobre HTTPS con TLS 1.2 o superior. Las conexiones HTTP deben redirigirse a HTTPS.'],
    ['RS-02', 'Autenticación MFA obligatoria', 'No se permite el acceso a ninguna operación funcional sin completar la autenticación de dos factores (LDAP + TOTP).'],
    ['RS-03', 'Rate limiting', 'Los endpoints de autenticación deben limitar las solicitudes a 5-10 por minuto por IP para prevenir ataques de fuerza bruta.'],
    ['RS-04', 'Caducidad de tokens', 'Los tokens JWT deben tener una caducidad máxima de 15 minutos. Los tickets OTP deben caducar en 5 minutos.'],
    ['RS-05', 'Headers de seguridad', 'Las respuestas HTTP deben incluir: HSTS, X-Content-Type-Options, X-Frame-Options, CSP, Referrer-Policy, Permissions-Policy.'],
    ['RS-06', 'Principio de mínimo privilegio', 'Los contenedores deben ejecutarse con usuarios no-root. La aplicación no debe tener acceso a recursos que no necesite.'],
    ['RS-07', 'Secretos seguros', 'Las claves criptográficas, contraseñas y tokens deben generarse con generadores criptográficamente seguros (CSPRNG).'],
    ['RS-08', 'Audit log inmutable', 'Todas las operaciones de seguridad deben registrarse en un log de auditoría separado con marca temporal UTC.'],
    ['RS-09', 'Firma digital obligatoria', 'Todo fichero subido debe ser firmado digitalmente con RSA-4096 PSS-SHA512. La clave privada debe tener permisos 0600.'],
    ['RS-10', 'Validación de entrada', 'Todos los datos de entrada deben validarse y sanitizarse para prevenir inyecciones (LDAP injection, path traversal).'],
]

add_table(['ID', 'Nombre', 'Descripción'], rs_data)

doc.add_heading('3.3 Requisitos de despliegue', level=2)

rd_data = [
    ['RD-01', 'Docker Compose', 'El sistema completo debe poder desplegarse con un único comando docker compose up -d.'],
    ['RD-02', 'Script de inicialización', 'Debe existir un script que genere automáticamente certificados TLS, secretos aleatorios, estructura de directorios y fichero .env.'],
    ['RD-03', 'Dependencias mínimas', 'El host solo debe requerir Docker Engine ≥ 20.10 y Docker Compose v2.'],
    ['RD-04', 'Recursos limitados', 'El sistema debe funcionar con un mínimo de 2 GB de RAM, distribuyendo los recursos entre los contenedores.'],
    ['RD-05', 'Volúmenes persistentes', 'Los datos (ficheros, claves RSA, logs, Redis, LDAP) deben almacenarse en volúmenes Docker para persistencia.'],
    ['RD-06', 'Red aislada', 'Los contenedores deben comunicarse a través de una red Docker privada (bridge) sin exponer puertos internos al host.'],
]

add_table(['ID', 'Nombre', 'Descripción'], rd_data)

doc.add_heading('3.4 Casos de uso', level=2)

doc.add_paragraph(
    'Los casos de uso principales del sistema se describen a continuación desde la '
    'perspectiva del actor "Usuario autenticado".'
)

doc.add_heading('CU-01: Autenticación completa (LDAP + MFA)', level=3)

doc.add_paragraph('Actor principal: Usuario')
doc.add_paragraph('Precondiciones: El usuario tiene credenciales LDAP válidas y una app de autenticación TOTP configurada.')
doc.add_paragraph('Flujo principal:')
cu01_steps = [
    'El usuario accede a la interfaz web del sistema.',
    'Introduce su UID y contraseña LDAP.',
    'El sistema valida las credenciales contra OpenLDAP.',
    'Si es la primera vez, se muestra un código QR para enrollment TOTP.',
    'El usuario introduce el código TOTP de 6 dígitos de su app de autenticación.',
    'El sistema verifica el código TOTP contra el secreto almacenado en LDAP.',
    'Se emite un token JWT con mfa_verified=true y se redirige al dashboard.',
]
for i, step in enumerate(cu01_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('Flujo alternativo 3a: Credenciales LDAP inválidas → Se muestra error, se registra en audit log.')
doc.add_paragraph('Flujo alternativo 6a: Código TOTP inválido → Se muestra error, se registra en audit log, se limpia el input.')
doc.add_paragraph('Postcondiciones: El usuario tiene un token JWT válido durante 15 minutos.')

doc.add_heading('CU-02: Subida de fichero con ticket OTP', level=3)

doc.add_paragraph('Actor principal: Usuario autenticado con MFA')
doc.add_paragraph('Precondiciones: Token JWT válido con mfa_verified=true.')
doc.add_paragraph('Flujo principal:')
cu02_steps = [
    'El usuario selecciona "Subir fichero" en el dashboard.',
    'Solicita un ticket de subida (POST /api/tickets/create con operation=upload).',
    'El sistema genera un ticket OTP, lo almacena en Redis con TTL=300s.',
    'Se muestra el ticket al usuario con un contador regresivo de 5 minutos.',
    'El usuario selecciona un fichero (drag & drop o selector) y pulsa "Subir".',
    'El sistema valida el ticket OTP (existencia, no usado, operación correcta).',
    'El fichero se almacena en /data/sftp-storage/{uid}/{filename}.',
    'Se calcula el hash SHA-512 del fichero.',
    'Se firma el hash con la clave RSA-4096 (PSS-SHA512).',
    'Se genera un recibo digital con hash, firma, timestamp y metadatos.',
    'Se almacenan los metadatos en Redis (hash, firma, fecha, receipt_id).',
    'Se envía un email al usuario con el recibo digital.',
    'Se muestra el recibo en la interfaz con log de terminal.',
]
for i, step in enumerate(cu02_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('Flujo alternativo 6a: Ticket inválido/expirado/ya usado → Se bloquea la subida, se registra en audit log.')
doc.add_paragraph('Postcondiciones: Fichero almacenado, firmado y con recibo digital generado.')

doc.add_heading('CU-03: Descarga de fichero con verificación de integridad', level=3)

doc.add_paragraph('Actor principal: Usuario autenticado con MFA')
doc.add_paragraph('Precondiciones: Token JWT válido, fichero previamente subido, hash SHA-512 del recibo de subida.')
doc.add_paragraph('Flujo principal:')
cu03_steps = [
    'El usuario selecciona "Descargar fichero" en el dashboard.',
    'Introduce el nombre del fichero y el hash SHA-512 del recibo de subida.',
    'El sistema genera un ticket OTP de descarga.',
    'Se valida el ticket y se localiza el fichero en disco.',
    'Se recalcula el hash SHA-512 del fichero actual en disco.',
    'Se compara el hash actual con el hash proporcionado por el usuario.',
    'Si los hashes coinciden, se verifica también la firma RSA almacenada.',
    'Si todo es correcto, se sirve el fichero para descarga.',
]
for i, step in enumerate(cu03_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph('Flujo alternativo 6a: Hash no coincide → FICHERO COMPROMETIDO. Se bloquea la descarga. Se abre un incidente automático. Se registra en audit log con nivel CRITICAL.')
doc.add_paragraph('Flujo alternativo 7a: Firma RSA no válida → Se bloquea la descarga por posible manipulación.')
doc.add_paragraph('Postcondiciones: Fichero descargado con integridad verificada, o incidente abierto.')

doc.add_heading('CU-04: Verificación manual de integridad', level=3)

doc.add_paragraph('Actor principal: Usuario autenticado o auditor')
doc.add_paragraph('Flujo: El usuario introduce nombre de fichero y hash SHA-512. El sistema compara el hash actual del fichero, verifica la firma RSA y retorna un resultado detallado indicando si la integridad está verificada o comprometida.')

page_break()

# ═══════════════════════════════════════════════════════════════
#  4. DISEÑO DE LA ARQUITECTURA
# ═══════════════════════════════════════════════════════════════

doc.add_heading('4. Diseño de la arquitectura', level=1)

doc.add_heading('4.1 Arquitectura general del sistema', level=2)

doc.add_paragraph(
    'El sistema sigue una arquitectura de microservicios ligera, orquestada mediante Docker '
    'Compose, donde cada componente funcional se ejecuta en un contenedor independiente con '
    'responsabilidades claramente definidas. La comunicación entre componentes se realiza '
    'exclusivamente a través de una red Docker privada (bridge) en la subred 172.28.0.0/16, '
    'sin exponer servicios internos al exterior.'
)

doc.add_paragraph('La arquitectura consta de los siguientes componentes:')

arch_components = [
    ['Nginx', 'Reverse proxy + terminación TLS', 'nginx:1.27-alpine', '443, 80', 'Punto de entrada único. Redirige HTTP→HTTPS. Aplica rate limiting, security headers y CSP. Proxifica al backend Flask.'],
    ['Flask App', 'Backend API REST', 'python:3.12-slim + Gunicorn', '5000 (interno)', 'API REST, lógica de negocio, autenticación, tickets, firma digital, gestión de ficheros. Sirve también el frontend estático.'],
    ['OpenLDAP', 'Directorio de usuarios', 'osixia/openldap:1.5.0', '389 (interno)', 'Almacena identidades de usuario, credenciales y secretos TOTP. Autenticación mediante LDAP bind.'],
    ['Redis', 'Almacén de tickets/sesiones', 'redis:7-alpine', '6379 (interno)', 'Tickets OTP, blacklist JWT, metadatos de ficheros, recibos digitales, incidentes de seguridad. Con AOF habilitado para persistencia.'],
    ['MailHog', 'Servidor SMTP (dev)', 'mailhog:v1.0.1', '8025 (Web UI)', 'Captura emails en desarrollo. En producción se reemplaza por relay SMTP real.'],
]

add_table(
    ['Componente', 'Función', 'Imagen Docker', 'Puertos', 'Descripción'],
    arch_components
)

doc.add_paragraph(
    'El diagrama de la arquitectura muestra el flujo de comunicación entre componentes:'
)

doc.add_paragraph(
    '[Usuario/Navegador] ──HTTPS──► [Nginx :443] ──HTTP──► [Flask :5000]\n'
    '                                                         │\n'
    '                                                    ┌────┴────┐\n'
    '                                               [Redis :6379] [LDAP :389]\n'
    '                                                    │\n'
    '                                              [MailHog :1025]'
)

doc.add_paragraph(
    'Flujo de una solicitud típica:'
)

flow_steps = [
    'El navegador del usuario envía una solicitud HTTPS al puerto 443.',
    'Nginx termina TLS, aplica rate limiting y security headers, y proxifica la solicitud al backend Flask en el puerto 5000.',
    'Flask procesa la solicitud: valida el JWT, consulta Redis para tickets/sesiones, autentica contra LDAP si es necesario.',
    'Flask almacena o recupera ficheros del volumen sftp-storage.',
    'Flask envía notificaciones por email a través de MailHog (SMTP puerto 1025).',
    'La respuesta recorre el camino inverso hasta el navegador.',
]
for i, step in enumerate(flow_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('4.2 Stack tecnológico y justificación', level=2)

stack_data = [
    ['Python 3.12', 'Lenguaje backend', 'Ecosistema maduro para criptografía (cryptography, PyJWT), LDAP (ldap3), web (Flask). Excelente soporte para operaciones de seguridad.'],
    ['Flask 3.0', 'Framework web', 'Microframework ligero, sin overhead innecesario. Ideal para APIs REST donde se necesita control fino sobre cada aspecto de la aplicación.'],
    ['Gunicorn 22.0', 'Servidor WSGI', 'Servidor de producción con soporte para múltiples workers (4 workers × 2 threads). Pre-fork model para concurrencia.'],
    ['cryptography 43.0', 'Biblioteca criptográfica', 'Wrapper de OpenSSL mantenido activamente. Soporte completo para RSA, PSS, SHA-512 y gestión de claves.'],
    ['PyJWT 2.9', 'Tokens JWT', 'Implementación estándar de JSON Web Tokens con soporte para HS256.'],
    ['ldap3 2.9', 'Cliente LDAP', 'Biblioteca Python pura para comunicación LDAP. No requiere dependencias de sistema adicionales.'],
    ['pyotp 2.9', 'TOTP/HOTP', 'Implementación del RFC 6238 (TOTP) y RFC 4226 (HOTP). Compatible con Google Authenticator.'],
    ['Redis 7', 'Almacén en memoria', 'Almacén clave-valor con TTL nativo. Ideal para tickets temporales, sesiones y blacklists.'],
    ['OpenLDAP 1.5', 'Directorio LDAP', 'Implementación open-source del protocolo LDAP. Directorio estándar para gestión de identidades.'],
    ['Nginx 1.27', 'Reverse proxy', 'Servidor web de alto rendimiento. Excelente soporte TLS, rate limiting y security headers.'],
    ['Docker Compose', 'Orquestación', 'Definición declarativa de la infraestructura. Despliegue reproducible con un solo comando.'],
    ['HTML/CSS/JS', 'Frontend SPA', 'Interfaz vanilla sin frameworks. Reduce la superficie de ataque (sin node_modules) y minimiza dependencias.'],
]

add_table(['Tecnología', 'Rol', 'Justificación'], stack_data)

doc.add_heading('4.3 Modelo de autenticación (LDAP + TOTP)', level=2)

doc.add_paragraph(
    'El modelo de autenticación implementa un flujo secuencial de dos fases que debe '
    'completarse en su totalidad antes de obtener acceso al sistema:'
)

doc.add_paragraph(
    'Fase 1 — Autenticación LDAP: El usuario envía sus credenciales (uid + contraseña) al '
    'endpoint POST /api/auth/login. El backend realiza un LDAP bind directo contra el '
    'servidor OpenLDAP, intentando autenticarse con el DN del usuario '
    '(uid={username},ou=users,dc=sftp,dc=secure,dc=local). Si el bind es exitoso, las '
    'credenciales son válidas y se recuperan los atributos del usuario (cn, mail, uid). '
    'En este punto se emite un token JWT "pre-MFA" con mfa_verified=false, que solo permite '
    'acceder al endpoint de verificación MFA.'
)

doc.add_paragraph(
    'Fase 2 — Verificación TOTP: El usuario envía el código TOTP de 6 dígitos junto con '
    'el token pre-MFA al endpoint POST /api/auth/verify-mfa. El backend decodifica el token '
    'pre-MFA, recupera el secreto TOTP del usuario desde LDAP (almacenado en el campo '
    'description con prefijo TOTP:), genera el código TOTP esperado para el momento actual '
    '(con ventana ±1) y lo compara con el proporcionado. Si coincide, se emite un token JWT '
    'completo con mfa_verified=true, que permite acceder a todas las funcionalidades.'
)

doc.add_paragraph(
    'Gestión de secretos TOTP: El secreto TOTP se genera la primera vez con '
    'pyotp.random_base32() (160 bits de entropía) y se almacena en LDAP en el campo '
    'description del usuario con formato "TOTP:{base32_secret}". Esta decisión de diseño '
    'centraliza la gestión de secretos en el directorio LDAP, evitando mantener un almacén '
    'separado. En un entorno de producción, se recomienda usar un campo LDAP personalizado '
    'o un atributo de esquema extendido.'
)

doc.add_heading('4.4 Modelo de autorización basado en tickets OTP', level=2)

doc.add_paragraph(
    'La autorización de operaciones de ficheros se basa en un sistema de tickets de un solo '
    'uso (OTP) que complementa la autenticación JWT. El objetivo es implementar el principio '
    'de "defense in depth": incluso con un JWT válido, el usuario necesita obtener un ticket '
    'específico para cada operación individual.'
)

doc.add_paragraph('Estructura de un ticket OTP en Redis:')

add_code_block(
    'Key:   ticket:{token_urlsafe_48bytes}\n'
    'Type:  Hash\n'
    'TTL:   300 segundos (configurable)\n'
    'Fields:\n'
    '  ticket_id:   Token criptográficamente aleatorio (48 bytes, URL-safe)\n'
    '  user_uid:    UID del usuario que solicitó el ticket\n'
    '  operation:   "upload" | "download"\n'
    '  filename:    Nombre del fichero (solo para download)\n'
    '  created_at:  Timestamp ISO 8601 UTC\n'
    '  used:        "false" | "true"\n'
    '  ip_address:  IP del solicitante'
)

doc.add_paragraph('Ciclo de vida del ticket:')

ticket_lifecycle = [
    'Creación: El usuario autenticado solicita un ticket indicando la operación (upload/download). Se genera un token aleatorio con secrets.token_urlsafe(48) y se almacena como hash en Redis con TTL=300s.',
    'Validación: Al recibir una solicitud de subida/descarga, el sistema busca el ticket en Redis, verifica que no haya sido usado (used=false), que la operación coincida, y que exista (no haya expirado).',
    'Consumo: El ticket se marca como used=true de forma atómica con HSET. El TTL se reduce a 60 segundos para mantener un breve registro de auditoría.',
    'Expiración: Redis elimina automáticamente el ticket tras el TTL sin intervención del backend.',
]
for i, step in enumerate(ticket_lifecycle, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('4.5 Modelo de firma digital y registro de integridad', level=2)

doc.add_paragraph(
    'El modelo de firma digital implementa un registro de integridad completo para cada '
    'fichero subido al sistema. El flujo criptográfico es:'
)

doc.add_paragraph(
    '1. Hash: Se calcula el hash SHA-512 del contenido completo del fichero, leyendo en '
    'bloques de 8192 bytes para soportar ficheros grandes sin cargar todo en memoria.\n\n'
    '2. Firma: Se firma el string hexadecimal del hash (no los bytes del fichero directamente) '
    'con RSA-4096 usando el esquema PSS (Probabilistic Signature Scheme) con SHA-512 como '
    'función hash interna y salt de longitud máxima.\n\n'
    '3. Recibo: Se genera un recibo digital estructurado (JSON) que incluye: receipt_id '
    '(UUID v4), timestamp (ISO 8601 UTC), operación, usuario, nombre de fichero, algoritmo '
    'de hash, hash del fichero, algoritmo de firma, firma RSA en base64, clave RSA utilizada '
    '(tamaño) y estado de integridad.\n\n'
    '4. Almacenamiento: El recibo completo se almacena en Redis como JSON serializado '
    '(key: receipt:{uid}:{filename}). Los metadatos de firma se almacenan por separado '
    'como hash de Redis (key: filemeta:{uid}:{filename}) para consulta rápida.'
)

doc.add_paragraph('Diagrama del flujo de firma:')

add_code_block(
    'Fichero → [SHA-512] → hash_hex → [RSA-4096 PSS-SHA512] → signature_b64\n'
    '                           │                                     │\n'
    '                           └──────── Recibo Digital ──────────────┘\n'
    '                                         │\n'
    '                              ┌──────────┼──────────┐\n'
    '                         [Redis]    [Email SMTP]  [API Response]'
)

doc.add_heading('4.6 Diagrama de flujo de operaciones', level=2)

doc.add_paragraph('Flujo completo de una operación de subida de fichero:')

add_code_block(
    '┌─────────────┐    ┌──────────────┐    ┌─────────────┐\n'
    '│  Usuario     │    │  Nginx       │    │  Flask API  │\n'
    '│  (Browser)   │    │  (TLS Proxy) │    │  (Backend)  │\n'
    '└──────┬───────┘    └──────┬───────┘    └──────┬──────┘\n'
    '       │                   │                   │\n'
    '  1. POST /api/auth/login  │                   │\n'
    '       │──────────────────►│──────────────────►│\n'
    '       │                   │     LDAP bind     │──► [OpenLDAP]\n'
    '       │                   │     QR + pre-JWT  │\n'
    '       │◄──────────────────│◄──────────────────│\n'
    '       │                   │                   │\n'
    '  2. POST /api/auth/verify-mfa                 │\n'
    '       │──────────────────►│──────────────────►│\n'
    '       │                   │   Verify TOTP     │──► [LDAP]\n'
    '       │                   │   Full JWT        │\n'
    '       │◄──────────────────│◄──────────────────│\n'
    '       │                   │                   │\n'
    '  3. POST /api/tickets/create                  │\n'
    '       │──────────────────►│──────────────────►│\n'
    '       │                   │  Store ticket     │──► [Redis]\n'
    '       │                   │  ticket_id        │\n'
    '       │◄──────────────────│◄──────────────────│\n'
    '       │                   │                   │\n'
    '  4. POST /api/sftp/upload (multipart)         │\n'
    '       │──────────────────►│──────────────────►│\n'
    '       │                   │  Validate ticket  │──► [Redis]\n'
    '       │                   │  Save file        │──► [Disk]\n'
    '       │                   │  SHA-512 + RSA    │\n'
    '       │                   │  Store meta       │──► [Redis]\n'
    '       │                   │  Send email       │──► [SMTP]\n'
    '       │                   │  Receipt          │\n'
    '       │◄──────────────────│◄──────────────────│\n'
)

doc.add_heading('4.7 Modelo de datos (Redis, LDAP)', level=2)

doc.add_paragraph('El sistema utiliza Redis como almacén principal de datos operativos y LDAP como directorio de identidades.')

doc.add_heading('Estructura de datos en Redis', level=3)

redis_keys = [
    ['ticket:{id}', 'Hash', '300s', 'Ticket OTP con metadatos de operación'],
    ['blacklist:{jti}', 'String', 'JWT_EXPIRY', 'JTI de tokens JWT invalidados (logout)'],
    ['receipt:{uid}:{filename}', 'String (JSON)', 'Sin TTL', 'Recibo digital completo serializado'],
    ['filemeta:{uid}:{filename}', 'Hash', 'Sin TTL', 'Metadatos del fichero: hash, firma, fecha, receipt_id, tamaño'],
    ['incident:{uuid}', 'Hash', 'Sin TTL', 'Incidente de seguridad: tipo, usuario, fichero, hashes, IP, timestamp, estado'],
]

add_table(['Clave', 'Tipo Redis', 'TTL', 'Descripción'], redis_keys)

doc.add_heading('Estructura de datos en LDAP', level=3)

doc.add_paragraph(
    'El directorio LDAP utiliza el esquema estándar inetOrgPerson con la siguiente '
    'estructura de árbol:'
)

add_code_block(
    'dc=sftp,dc=secure,dc=local          (dominio base)\n'
    '└── ou=users                         (unidad organizativa)\n'
    '    ├── uid=admin                    (Administrador)\n'
    '    │   ├── cn: Administrador SFTP\n'
    '    │   ├── sn: Admin\n'
    '    │   ├── mail: admin@secure.local\n'
    '    │   ├── userPassword: {hash}\n'
    '    │   ├── employeeNumber: 001\n'
    '    │   └── description: TOTP:{base32_secret}\n'
    '    ├── uid=operador\n'
    '    └── uid=auditor'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  5. IMPLEMENTACIÓN
# ═══════════════════════════════════════════════════════════════

doc.add_heading('5. Implementación', level=1)

doc.add_paragraph(
    'En este capítulo se describe la implementación detallada de cada componente del sistema, '
    'incluyendo fragmentos de código significativos, decisiones de diseño y configuraciones '
    'relevantes.'
)

doc.add_heading('5.1 Contenedorización con Docker', level=2)

doc.add_paragraph(
    'El sistema se despliega mediante Docker Compose con cinco servicios. El fichero '
    'docker-compose.yml define la orquestación completa, incluyendo dependencias entre '
    'servicios (depends_on con healthcheck), límites de recursos (memory, cpus), volúmenes '
    'persistentes y configuración de red.'
)

doc.add_paragraph('Dockerfile de la aplicación (docker/Dockerfile.app):')

add_code_block(
    'FROM python:3.12-slim\n\n'
    '# Seguridad: usuario no-root\n'
    'RUN groupadd -r sftpapp && useradd -r -g sftpapp -d /app -s /sbin/nologin sftpapp\n\n'
    '# Dependencias del sistema\n'
    'RUN apt-get update && apt-get install -y --no-install-recommends \\\n'
    '    curl libldap2-dev libsasl2-dev \\\n'
    '    && rm -rf /var/lib/apt/lists/*\n\n'
    '# Directorios con permisos correctos\n'
    'RUN mkdir -p /data/sftp-storage /data/keys /data/logs /app \\\n'
    '    && chown -R sftpapp:sftpapp /data /app\n\n'
    'WORKDIR /app\n'
    'COPY backend/requirements.txt .\n'
    'RUN pip install --no-cache-dir -r requirements.txt\n'
    'COPY backend/ ./backend/\n'
    'COPY frontend/ ./frontend/\n'
    'RUN chown -R sftpapp:sftpapp /app\n'
    'USER sftpapp\n\n'
    'EXPOSE 5000\n'
    'HEALTHCHECK --interval=30s --timeout=10s --retries=3 \\\n'
    '    CMD curl -f http://localhost:5000/api/health || exit 1\n\n'
    'CMD ["gunicorn", "--bind", "0.0.0.0:5000", "--workers", "4",\n'
    '     "--threads", "2", "--timeout", "120", "--chdir", "/app/backend", "app:app"]'
)

doc.add_paragraph(
    'Decisiones de seguridad en el Dockerfile:'
)

docker_security = [
    'Imagen base slim: python:3.12-slim minimiza la superficie de ataque eliminando paquetes innecesarios.',
    'Usuario no-root: la aplicación se ejecuta como sftpapp (UID no privilegiado) con shell /sbin/nologin para prevenir login interactivo.',
    'Limpieza de caché: Se elimina la caché de apt y pip para reducir el tamaño de la imagen y no dejar paquetes de compilación.',
    'Healthcheck nativo: Docker monitoriza la salud del contenedor y puede reiniciarlo automáticamente.',
    'Gunicorn como servidor: 4 workers con 2 threads cada uno proporcionan concurrencia sin los riesgos de seguridad de un servidor de desarrollo.',
]
for item in docker_security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Volúmenes Docker definidos:')

volumes_data = [
    ['sftp-storage', '/data/sftp-storage', 'Ficheros subidos por los usuarios, organizados por UID'],
    ['signing-keys', '/data/keys', 'Clave privada RSA-4096 para firma digital (signing_key.pem)'],
    ['app-logs', '/data/logs', 'Logs de la aplicación: sftp-service.log, audit.log, access.log'],
    ['ldap-data', '/var/lib/ldap', 'Base de datos del directorio LDAP'],
    ['ldap-config', '/etc/ldap/slapd.d', 'Configuración de slapd (LDAP server)'],
    ['redis-data', '/data', 'Datos persistentes de Redis (AOF + RDB)'],
]

add_table(['Volumen', 'Mount point', 'Contenido'], volumes_data)

doc.add_heading('5.2 Backend Flask: API REST', level=2)

doc.add_paragraph(
    'El backend del sistema está implementado en un único fichero Python (backend/app.py) '
    'que contiene la configuración de Flask, la lógica de negocio y la definición de todos '
    'los endpoints de la API REST. Esta decisión de diseño (monolito funcional) se justifica '
    'por la cohesión alta entre los componentes: la autenticación, los tickets y la firma '
    'digital están intrínsecamente relacionados y comparten estado.'
)

doc.add_paragraph('Endpoints de la API REST:')

api_data = [
    ['POST', '/api/auth/login', '10/min', 'No', 'Autenticación LDAP (Fase 1 MFA)'],
    ['POST', '/api/auth/verify-mfa', '5/min', 'Pre-MFA', 'Verificación TOTP (Fase 2 MFA)'],
    ['POST', '/api/auth/logout', '100/min', 'Sí', 'Invalidación de token JWT (blacklist)'],
    ['POST', '/api/tickets/create', '100/min', 'Sí', 'Generación de ticket OTP'],
    ['POST', '/api/sftp/upload', '20/min', 'Ticket', 'Subida de fichero con firma digital'],
    ['POST', '/api/sftp/download', '20/min', 'Ticket', 'Descarga con verificación de integridad'],
    ['GET', '/api/sftp/files', '100/min', 'Sí', 'Listado de ficheros del usuario'],
    ['POST', '/api/sftp/verify-receipt', '100/min', 'Sí', 'Verificación manual de integridad'],
    ['GET', '/api/incidents', '100/min', 'Sí', 'Listado de incidentes de seguridad'],
    ['GET', '/api/health', 'Sin límite', 'No', 'Health check (Redis + LDAP)'],
]

add_table(['Método', 'Ruta', 'Rate Limit', 'Auth', 'Descripción'], api_data)

doc.add_paragraph(
    'Configuración de la aplicación: Todas las variables de configuración se cargan desde '
    'variables de entorno con valores por defecto seguros. Los secretos (SECRET_KEY, '
    'JWT_SECRET) se generan con secrets.token_hex(64) si no se proporcionan, garantizando '
    '512 bits de entropía. Los parámetros de conexión a servicios (LDAP, Redis, SMTP) se '
    'configuran mediante variables de entorno inyectadas por Docker Compose.'
)

doc.add_paragraph(
    'Middleware de autenticación: El decorador @require_auth implementa la verificación '
    'de JWT en cada endpoint protegido. Comprueba: presencia del token en el header '
    'Authorization (formato "Bearer {token}"), validez del token (firma HS256, no expirado), '
    'campo mfa_verified=true, y que el JTI (JWT ID) no esté en la blacklist de Redis. Si '
    'alguna verificación falla, se retorna un error 401 o 403 con mensaje descriptivo.'
)

doc.add_heading('5.3 Servicio de autenticación LDAP', level=2)

doc.add_paragraph(
    'La integración LDAP utiliza la biblioteca ldap3 (Python pura) para comunicarse con '
    'el servidor OpenLDAP. Se implementan dos funciones principales:'
)

doc.add_paragraph(
    'ldap_authenticate(username, password): Realiza un LDAP bind directo con el DN del '
    'usuario para validar las credenciales. Si el bind es exitoso, busca los atributos del '
    'usuario (cn, mail, uid) y los retorna como diccionario. El uso de bind directo '
    '(en lugar de buscar primero y comparar contraseñas) es la práctica recomendada de '
    'seguridad, ya que la contraseña nunca se transfiere en claro ni se almacena en '
    'variables intermedias del backend.'
)

doc.add_paragraph(
    'ldap_get_totp_secret(username): Utiliza una conexión administrativa (cn=admin) para '
    'leer y modificar el campo description del usuario, donde se almacena el secreto TOTP. '
    'Si el usuario no tiene un secreto TOTP configurado, se genera uno nuevo con '
    'pyotp.random_base32() y se almacena en LDAP con el prefijo "TOTP:".'
)

doc.add_paragraph('Configuración LDAP del sistema:')

add_code_block(
    'LDAP_URI          = ldap://openldap:389\n'
    'LDAP_BASE_DN      = dc=sftp,dc=secure,dc=local\n'
    'LDAP_USERS_DN     = ou=users,dc=sftp,dc=secure,dc=local\n'
    'LDAP_ADMIN_DN     = cn=admin,dc=sftp,dc=secure,dc=local\n'
    'LDAP_ADMIN_PASSWORD = {generado por init.sh}'
)

doc.add_paragraph(
    'El servidor OpenLDAP se inicializa con un fichero bootstrap.ldif que crea la unidad '
    'organizativa ou=users y tres usuarios de prueba (admin, operador, auditor) con '
    'contraseñas predefinidas. En un entorno de producción, el directorio LDAP se integraría '
    'con el servicio de identidades corporativo existente.'
)

doc.add_heading('5.4 Generación de QR y verificación TOTP', level=2)

doc.add_paragraph(
    'La generación de códigos QR para enrollment TOTP se implementa con las bibliotecas '
    'pyotp (generación del URI otpauth://) y qrcode (generación de la imagen QR):'
)

add_code_block(
    'def generate_totp_qr(username, secret):\n'
    '    totp = pyotp.TOTP(secret)\n'
    '    provisioning_uri = totp.provisioning_uri(\n'
    '        name=username,\n'
    '        issuer_name="SFTP Secure Service"\n'
    '    )\n'
    '    qr = qrcode.QRCode(version=1, box_size=6, border=2)\n'
    '    qr.add_data(provisioning_uri)\n'
    '    qr.make(fit=True)\n'
    '    img = qr.make_image(fill_color="black", back_color="white")\n'
    '    buffer = BytesIO()\n'
    '    img.save(buffer, format="PNG")\n'
    '    buffer.seek(0)\n'
    '    return base64.b64encode(buffer.getvalue()).decode("utf-8")'
)

doc.add_paragraph(
    'El URI de provisioning sigue el formato estándar: '
    'otpauth://totp/SFTP%20Secure%20Service:{username}?secret={base32}&issuer=SFTP%20Secure%20Service'
)

doc.add_paragraph(
    'La verificación TOTP acepta una ventana de tolerancia de ±1 período (valid_window=1), '
    'lo que significa que se aceptan códigos del período actual, del período anterior y del '
    'siguiente. Esto compensa posibles desincronizaciones de reloj entre el servidor y el '
    'dispositivo del usuario de hasta ±30 segundos.'
)

doc.add_heading('5.5 Gestión de tickets OTP en Redis', level=2)

doc.add_paragraph(
    'Los tickets OTP se almacenan como hashes de Redis, aprovechando las capacidades nativas '
    'de TTL de Redis para la caducidad automática. Se implementan dos funciones principales:'
)

add_code_block(
    'def create_ticket(user_uid, operation, filename=None):\n'
    '    r = get_redis()\n'
    '    ticket_id = secrets.token_urlsafe(48)  # 384 bits de entropía\n'
    '    ticket_data = {\n'
    '        "ticket_id": ticket_id,\n'
    '        "user_uid": user_uid,\n'
    '        "operation": operation,     # "upload" | "download"\n'
    '        "filename": filename or "", \n'
    '        "created_at": datetime.now(timezone.utc).isoformat(),\n'
    '        "used": "false",\n'
    '        "ip_address": request.remote_addr,\n'
    '    }\n'
    '    r.hset(f"ticket:{ticket_id}", mapping=ticket_data)\n'
    '    r.expire(f"ticket:{ticket_id}", TICKET_TTL_SECONDS)  # 300s\n'
    '    return ticket_id'
)

doc.add_paragraph(
    'La validación del ticket implementa tres comprobaciones atómicas: existencia '
    '(el ticket no ha expirado), estado de uso (used=false), y coincidencia de operación '
    '(upload/download). Al consumir el ticket, se marca como used=true y se reduce el TTL '
    'a 60 segundos para mantener un breve registro de auditoría.'
)

doc.add_heading('5.6 Firma digital RSA-4096 (PSS-SHA512)', level=2)

doc.add_paragraph(
    'La gestión de claves RSA implementa un patrón de carga perezosa (lazy loading) con '
    'generación automática si la clave no existe:'
)

add_code_block(
    'def load_or_create_signing_key():\n'
    '    key_path = Path(RSA_KEY_PATH)\n'
    '    if key_path.exists():\n'
    '        with open(key_path, "rb") as f:\n'
    '            private_key = serialization.load_pem_private_key(\n'
    '                f.read(), password=None, backend=default_backend()\n'
    '            )\n'
    '    else:\n'
    '        private_key = rsa.generate_private_key(\n'
    '            public_exponent=65537,\n'
    '            key_size=4096,\n'
    '            backend=default_backend()\n'
    '        )\n'
    '        with open(key_path, "wb") as f:\n'
    '            f.write(private_key.private_bytes(\n'
    '                encoding=serialization.Encoding.PEM,\n'
    '                format=serialization.PrivateFormat.PKCS8,\n'
    '                encryption_algorithm=serialization.NoEncryption()\n'
    '            ))\n'
    '        os.chmod(key_path, 0o600)  # Solo lectura para owner\n'
    '    return private_key'
)

doc.add_paragraph(
    'La función de firma utiliza PSS con salt de longitud máxima, lo que maximiza la '
    'seguridad del esquema:'
)

add_code_block(
    'def sign_hash(file_hash: str) -> str:\n'
    '    key = get_signing_key()\n'
    '    signature = key.sign(\n'
    '        file_hash.encode("utf-8"),\n'
    '        padding.PSS(\n'
    '            mgf=padding.MGF1(hashes.SHA512()),\n'
    '            salt_length=padding.PSS.MAX_LENGTH\n'
    '        ),\n'
    '        hashes.SHA512()\n'
    '    )\n'
    '    return base64.b64encode(signature).decode("utf-8")'
)

doc.add_heading('5.7 Verificación de integridad en descarga', level=2)

doc.add_paragraph(
    'El mecanismo de verificación de integridad en la descarga es uno de los elementos '
    'diferenciadores del sistema. A diferencia de los protocolos SFTP/FTPS que solo '
    'verifican la integridad del canal de transmisión, este sistema verifica la integridad '
    'del fichero almacenado, detectando cualquier modificación posterior a la subida.'
)

doc.add_paragraph('El flujo de verificación en la descarga (/api/sftp/download) es:')

verify_steps = [
    'El usuario proporciona el ticket OTP, el nombre del fichero y el hash SHA-512 original (obtenido del recibo de subida).',
    'Se valida y consume el ticket OTP.',
    'Se localiza el fichero en disco (/data/sftp-storage/{uid}/{filename}).',
    'Se recalcula el hash SHA-512 del fichero actual en disco.',
    'Se compara el hash actual con el hash proporcionado por el usuario.',
    'Si los hashes NO coinciden: se registra un evento CRITICAL en el audit log, se crea un incidente de seguridad en Redis con tipo FILE_INTEGRITY_FAILURE, se retorna un error HTTP 409 (Conflict) con detalles del incidente.',
    'Si los hashes coinciden: se recupera la firma RSA almacenada en Redis y se verifica contra el hash actual. Si la firma no es válida, se bloquea la descarga por posible manipulación.',
    'Si todo es correcto: se sirve el fichero con send_file().',
]
for i, step in enumerate(verify_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('5.8 Notificaciones por correo electrónico', level=2)

doc.add_paragraph(
    'El sistema envía notificaciones por correo electrónico tras cada subida exitosa de '
    'fichero. El email contiene el recibo digital completo en dos formatos: texto plano '
    '(para compatibilidad con clientes de correo ligeros y archivado) y HTML (con diseño '
    'estilizado coherente con la interfaz del sistema).'
)

doc.add_paragraph(
    'El email HTML utiliza un diseño oscuro con acentos verdes que reproduce la estética '
    'de la interfaz terminal web del sistema. Incluye secciones diferenciadas para los '
    'detalles de la operación, la firma digital (con el hash SHA-512 resaltado en naranja '
    'y la firma RSA en azul) y una advertencia de seguridad instando al usuario a conservar '
    'el correo como comprobante.'
)

doc.add_paragraph(
    'La configuración SMTP se inyecta mediante variables de entorno. En desarrollo se '
    'utiliza MailHog (puerto 1025) como servidor SMTP sin autenticación que captura todos '
    'los correos y permite visualizarlos en su interfaz web (puerto 8025). En producción, '
    'se configuraría un relay SMTP autenticado con TLS (ej. Amazon SES, SendGrid o un '
    'servidor corporativo).'
)

doc.add_heading('5.9 Frontend: interfaz terminal web', level=2)

doc.add_paragraph(
    'La interfaz de usuario es una Single Page Application (SPA) implementada en un único '
    'fichero HTML (frontend/index.html) con HTML, CSS y JavaScript vanilla, sin dependencias '
    'de frameworks externos. Esta decisión de diseño tiene dos motivaciones:'
)

frontend_reasons = [
    'Seguridad: al no depender de npm, webpack u otros sistemas de build, se elimina completamente el riesgo de ataques a la cadena de suministro (supply chain attacks) a través de paquetes maliciosos en node_modules.',
    'Simplicidad operativa: el frontend se sirve como fichero estático desde el propio backend Flask, sin necesidad de un servidor adicional ni proceso de build.',
]
for item in frontend_reasons:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'La interfaz implementa un diseño de "terminal de seguridad" con fondo oscuro, '
    'tipografía monoespaciada (JetBrains Mono, IBM Plex Mono), acentos verdes sobre negro '
    'y efectos visuales como scan lines, pulsos de borde y animaciones de entrada. Este '
    'diseño no es meramente estético: transmite visualmente que el usuario está interactuando '
    'con un sistema de seguridad, reforzando la importancia de las operaciones que realiza.'
)

doc.add_paragraph('La SPA implementa las siguientes vistas:')

views_data = [
    ['viewLogin', 'Formulario de autenticación LDAP (Paso 1/2)'],
    ['viewMFA', 'Verificación TOTP con código QR y 6 dígitos (Paso 2/2)'],
    ['viewDashboard', 'Panel principal con 4 acciones: Subir, Descargar, Mis Ficheros, Verificar'],
    ['viewUpload', 'Flujo de subida: solicitar ticket → seleccionar fichero → recibo digital'],
    ['viewDownload', 'Flujo de descarga: nombre + hash → ticket → verificación → descarga'],
    ['viewFiles', 'Listado de ficheros del usuario con metadatos'],
    ['viewVerify', 'Verificación manual de integridad por hash'],
]

add_table(['Vista', 'Descripción'], views_data)

doc.add_paragraph(
    'Características del frontend:'
)

frontend_features = [
    'Input TOTP con 6 dígitos individuales, navegación automática entre campos, soporte para pegado (paste) de códigos completos.',
    'Drop zone para arrastrar ficheros con feedback visual (border color change, nombre del fichero seleccionado).',
    'Contador regresivo visual del TTL del ticket OTP en tiempo real.',
    'Terminal log que simula una consola de seguridad mostrando cada paso de la operación con timestamps.',
    'Health check automático cada 30 segundos con indicador visual de estado (punto verde/rojo).',
    'Animaciones de entrada (fadeIn) para transiciones suaves entre vistas.',
    'Alertas contextuales con colores diferenciados: verde (éxito), rojo (error), naranja (warning).',
]
for item in frontend_features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.10 Proxy inverso y terminación TLS (Nginx)', level=2)

doc.add_paragraph(
    'Nginx actúa como punto de entrada único al sistema, proporcionando terminación TLS, '
    'rate limiting, security headers y reverse proxy hacia el backend Flask.'
)

doc.add_paragraph('Configuración TLS:')

add_code_block(
    'ssl_protocols       TLSv1.2 TLSv1.3;\n'
    'ssl_ciphers         ECDHE-ECDSA-AES128-GCM-SHA256:\n'
    '                    ECDHE-RSA-AES128-GCM-SHA256:\n'
    '                    ECDHE-ECDSA-AES256-GCM-SHA384:\n'
    '                    ECDHE-RSA-AES256-GCM-SHA384;\n'
    'ssl_prefer_server_ciphers on;\n'
    'ssl_session_cache   shared:SSL:10m;\n'
    'ssl_session_timeout 10m;\n'
    'ssl_session_tickets off;'
)

doc.add_paragraph('Rate limiting por zonas:')

rate_limits = [
    ['auth', '5 req/min', 'burst=3', 'Endpoints de autenticación (/api/auth/*)'],
    ['api', '30 req/s', 'burst=20', 'Endpoints generales de la API (/api/*)'],
]

add_table(['Zona', 'Rate', 'Burst', 'Aplicación'], rate_limits)

doc.add_paragraph('Security headers configurados en Nginx:')

headers_data = [
    ['Strict-Transport-Security', 'max-age=63072000; includeSubDomains; preload', 'Fuerza HTTPS durante 2 años'],
    ['X-Content-Type-Options', 'nosniff', 'Previene MIME sniffing'],
    ['X-Frame-Options', 'DENY', 'Previene clickjacking'],
    ['X-XSS-Protection', '1; mode=block', 'Activa filtro XSS del navegador'],
    ['Content-Security-Policy', "default-src 'self'; script-src 'self' 'unsafe-inline'...", 'Restringe orígenes de contenido'],
    ['Referrer-Policy', 'strict-origin-when-cross-origin', 'Limita información del referer'],
    ['Permissions-Policy', 'camera=(), microphone=(), geolocation=()', 'Deshabilita APIs del navegador innecesarias'],
]

add_table(['Header', 'Valor', 'Propósito'], headers_data)

doc.add_paragraph(
    'El servidor Nginx también oculta la información de versión (server_tokens off) y '
    'configura un tamaño máximo de subida de 500 MB (client_max_body_size 500M) con '
    'timeouts extendidos para operaciones de upload (proxy_read_timeout 600s).'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  6. DESPLIEGUE SEGURO
# ═══════════════════════════════════════════════════════════════

doc.add_heading('6. Despliegue seguro', level=1)

doc.add_heading('6.1 Entorno de despliegue (Ubuntu Server)', level=2)

doc.add_paragraph(
    'El sistema está diseñado para desplegarse en un servidor Ubuntu Server 22.04 LTS o '
    'superior con los siguientes requisitos mínimos:'
)

deploy_reqs = [
    ['Sistema operativo', 'Ubuntu Server 22.04 LTS (o RHEL 8+, Debian 12+)'],
    ['CPU', '2 vCPUs (recomendado 4 para producción)'],
    ['RAM', '2 GB mínimo (recomendado 4 GB)'],
    ['Disco', '20 GB para SO + 50 GB para almacenamiento de ficheros'],
    ['Docker Engine', '≥ 20.10'],
    ['Docker Compose', 'v2 (integrado en Docker Engine moderno)'],
    ['Red', 'IP fija, acceso a puertos 80 y 443'],
]

add_table(['Requisito', 'Especificación'], deploy_reqs)

doc.add_paragraph('Proceso de despliegue:')

add_code_block(
    '# 1. Clonar repositorio\n'
    'git clone <repositorio> sftp-ticket-service\n'
    'cd sftp-ticket-service\n\n'
    '# 2. Ejecutar script de inicialización\n'
    'chmod +x scripts/init.sh\n'
    './scripts/init.sh --dev    # o --prod para producción\n\n'
    '# 3. Levantar servicios\n'
    'docker compose up -d\n\n'
    '# 4. Esperar a que LDAP esté disponible (~15s)\n'
    'sleep 15\n\n'
    '# 5. Crear usuarios de prueba en LDAP\n'
    'docker compose exec app bash /app/scripts/setup_ldap_users.sh\n\n'
    '# 6. Verificar salud del sistema\n'
    'curl -k https://localhost/api/health'
)

doc.add_heading('6.2 Configuración de red y firewalls', level=2)

doc.add_paragraph(
    'La configuración de red del sistema se basa en el principio de mínima exposición. '
    'Solo los puertos estrictamente necesarios se exponen al exterior:'
)

ports_data = [
    ['443', 'Nginx', 'HTTPS — Punto de entrada principal', 'Sí'],
    ['80', 'Nginx', 'HTTP — Solo para redirect a HTTPS', 'Sí'],
    ['8025', 'MailHog', 'Web UI (solo dev)', 'No (solo dev)'],
    ['389', 'OpenLDAP', 'LDAP (solo Docker network)', 'No'],
    ['6379', 'Redis', 'Redis (solo Docker network)', 'No'],
    ['5000', 'Flask', 'API (solo Docker network)', 'No'],
]

add_table(['Puerto', 'Servicio', 'Descripción', 'Expuesto al host'], ports_data)

doc.add_paragraph('Reglas de firewall recomendadas (ufw):')

add_code_block(
    '# Permitir solo SSH, HTTP y HTTPS\n'
    'ufw default deny incoming\n'
    'ufw default allow outgoing\n'
    'ufw allow 22/tcp    # SSH (administración)\n'
    'ufw allow 80/tcp    # HTTP (redirect)\n'
    'ufw allow 443/tcp   # HTTPS\n'
    'ufw enable'
)

doc.add_paragraph(
    'La red Docker interna (sftp-network) utiliza la subred 172.28.0.0/16, aislada del '
    'host y de otras redes Docker. Los servicios internos (Redis, LDAP, Flask) solo son '
    'accesibles desde otros contenedores de la misma red.'
)

doc.add_heading('6.3 Gestión de secretos y variables de entorno', level=2)

doc.add_paragraph(
    'Los secretos del sistema se gestionan mediante un fichero .env que se genera '
    'automáticamente con el script init.sh. El fichero tiene permisos 600 (solo lectura '
    'para el propietario) y está incluido en .gitignore para evitar su inclusión accidental '
    'en el repositorio.'
)

doc.add_paragraph('Variables de entorno con secretos:')

secrets_data = [
    ['SECRET_KEY', 'openssl rand -hex 64', '512 bits', 'Clave secreta de Flask para sesiones'],
    ['JWT_SECRET', 'openssl rand -hex 64', '512 bits', 'Secreto para firma de tokens JWT (HS256)'],
    ['LDAP_ADMIN_PASSWORD', 'openssl rand -base64 24', '~144 bits', 'Contraseña del admin de LDAP'],
    ['LDAP_CONFIG_PASSWORD', 'openssl rand -base64 24', '~144 bits', 'Contraseña de configuración de LDAP'],
    ['REDIS_PASSWORD', 'openssl rand -base64 24', '~144 bits', 'Contraseña de autenticación de Redis'],
]

add_table(['Variable', 'Generación', 'Entropía', 'Propósito'], secrets_data)

doc.add_paragraph(
    'En un entorno de producción más avanzado, se recomienda migrar la gestión de secretos '
    'a HashiCorp Vault, AWS Secrets Manager o Azure Key Vault, que proporcionan rotación '
    'automática, audit trail y control de acceso basado en roles.'
)

doc.add_heading('6.4 Certificados TLS', level=2)

doc.add_paragraph(
    'El script init.sh genera certificados TLS autofirmados para desarrollo con el '
    'siguiente comando:'
)

add_code_block(
    'openssl req -x509 -nodes -days 365 -newkey rsa:2048 \\\n'
    '    -keyout docker/nginx/ssl/server.key \\\n'
    '    -out docker/nginx/ssl/server.crt \\\n'
    '    -subj "/C=ES/ST=Madrid/L=Madrid/O=SFTP Secure Dev/CN=localhost"'
)

doc.add_paragraph(
    'Para producción, se deben reemplazar estos certificados por certificados emitidos '
    'por una Autoridad Certificadora reconocida. Las opciones recomendadas son:'
)

tls_options = [
    'Let\'s Encrypt: certificados gratuitos con renovación automática cada 90 días. Ideal para servicios públicos. Se puede integrar con certbot como sidecar container.',
    'CA corporativa: si la organización tiene una PKI interna, se solicitan los certificados a la CA corporativa.',
    'CA comercial: DigiCert, Sectigo, etc. Para servicios que requieran Extended Validation (EV).',
]
for item in tls_options:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.5 Bastionado de contenedores', level=2)

doc.add_paragraph(
    'Las medidas de bastionado (hardening) implementadas en los contenedores Docker son:'
)

hardening = [
    ['Usuario no-root', 'La aplicación Flask se ejecuta como usuario sftpapp sin privilegios de sistema ni shell interactiva.'],
    ['Imagen base slim', 'python:3.12-slim minimiza los paquetes instalados y la superficie de ataque.'],
    ['Límites de recursos', 'Cada contenedor tiene límites de memoria (128-512 MB) y CPU (0.5-1.0) para prevenir denegación de servicio.'],
    ['Read-only mounts', 'Los ficheros de configuración (nginx.conf, bootstrap.ldif) se montan como read-only (:ro).'],
    ['No capabilities escalation', 'Los contenedores no solicitan capabilities adicionales de Linux.'],
    ['Health checks', 'Cada servicio crítico tiene un healthcheck que permite a Docker reiniciarlo automáticamente.'],
    ['Restart policy', 'Los contenedores se configuran con restart: unless-stopped para recuperación automática tras fallos.'],
    ['Redis autenticado', 'Redis requiere contraseña para todas las conexiones (requirepass).'],
    ['Redis sin acceso externo', 'El puerto 6379 no se expone al host, solo es accesible desde la red Docker interna.'],
]

add_table(['Medida', 'Descripción'], hardening)

doc.add_heading('6.6 Logs y auditoría', level=2)

doc.add_paragraph(
    'El sistema implementa un modelo de logging en tres niveles:'
)

doc.add_paragraph(
    '1. Log de aplicación (sftp-service.log): Registro general de eventos de la aplicación '
    'con nivel INFO y superior. Incluye errores, warnings y eventos informativos.'
)

doc.add_paragraph(
    '2. Log de auditoría (audit.log): Registro dedicado de eventos de seguridad con formato '
    'estructurado. Cada entrada incluye timestamp UTC, tipo de evento y parámetros relevantes. '
    'Eventos registrados:'
)

audit_events = [
    ['LDAP_AUTH_SUCCESS', 'Autenticación LDAP exitosa', 'user, ip'],
    ['LDAP_AUTH_FAILURE', 'Autenticación LDAP fallida', 'user, ip, reason'],
    ['MFA_SUCCESS', 'Verificación TOTP exitosa', 'user, ip'],
    ['MFA_FAILURE', 'Verificación TOTP fallida', 'user, ip'],
    ['TICKET_CREATED', 'Ticket OTP generado', 'id, user, operation, filename, ip, ttl'],
    ['TICKET_CONSUMED', 'Ticket OTP utilizado', 'id, user, operation, ip'],
    ['TICKET_INVALID', 'Intento de uso de ticket inválido', 'id, reason, ip'],
    ['TICKET_REUSE_ATTEMPT', 'Intento de reutilización de ticket', 'id, user, ip'],
    ['UPLOAD_SUCCESS', 'Fichero subido correctamente', 'user, file, hash, receipt, ip'],
    ['UPLOAD_DENIED', 'Subida denegada', 'reason, ticket, ip'],
    ['DOWNLOAD_SUCCESS', 'Fichero descargado con verificación', 'user, file, ip'],
    ['DOWNLOAD_DENIED', 'Descarga denegada', 'reason, ticket, ip'],
    ['FILE_COMPROMISED', 'Fichero comprometido detectado (CRITICAL)', 'user, file, hashes, ip'],
    ['SIGNATURE_MISMATCH', 'Firma RSA no válida (CRITICAL)', 'user, file, ip'],
    ['LOGOUT', 'Cierre de sesión', 'user, ip'],
]

add_table(['Evento', 'Descripción', 'Parámetros'], audit_events)

doc.add_paragraph(
    '3. Log de acceso Nginx (access.log): Registro de todas las solicitudes HTTP con '
    'formato extendido que incluye IP real, request time, status code y user agent.'
)

doc.add_heading('6.7 Plan de mantenimiento y actualizaciones', level=2)

doc.add_paragraph('El plan de mantenimiento recomendado incluye las siguientes actividades periódicas:')

maintenance = [
    ['Diario', 'Verificación de health checks y revisión de logs de auditoría'],
    ['Semanal', 'Actualización de imágenes Docker base (docker compose pull)'],
    ['Mensual', 'Revisión de CVEs en dependencias Python (pip audit)'],
    ['Trimestral', 'Rotación de la clave RSA de firma digital y re-firma de ficheros existentes'],
    ['Semestral', 'Revisión de configuración TLS (ssllabs.com), renovación de certificados'],
    ['Anual', 'Auditoría de seguridad completa, test de penetración, revisión de políticas de acceso'],
]

add_table(['Frecuencia', 'Actividad'], maintenance)

doc.add_paragraph(
    'Procedimiento de actualización de la aplicación:'
)

add_code_block(
    '# 1. Pull de la última versión\n'
    'git pull origin main\n\n'
    '# 2. Rebuild de la imagen de la aplicación\n'
    'docker compose build app\n\n'
    '# 3. Actualización rolling (sin downtime)\n'
    'docker compose up -d --no-deps app\n\n'
    '# 4. Verificación de salud\n'
    'curl -k https://localhost/api/health\n\n'
    '# 5. Revisión de logs\n'
    'docker compose logs -f app --tail=50'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  7. GESTIÓN DE INCIDENTES
# ═══════════════════════════════════════════════════════════════

doc.add_heading('7. Gestión de incidentes', level=1)

doc.add_paragraph(
    'Este capítulo describe el modelo de detección y respuesta a incidentes de seguridad '
    'implementado en el sistema, alineado con el marco de referencia NIST SP 800-61 Rev. 2 '
    '(Computer Security Incident Handling Guide).'
)

doc.add_heading('7.1 Modelo de detección de ficheros comprometidos', level=2)

doc.add_paragraph(
    'El sistema implementa detección automática de ficheros comprometidos en dos puntos '
    'del flujo operativo:'
)

doc.add_paragraph(
    '1. En la descarga (/api/sftp/download): Cuando el usuario solicita descargar un '
    'fichero, el sistema recalcula el hash SHA-512 del fichero actual en disco y lo compara '
    'con dos referencias:'
)

doc.add_paragraph(
    '   a) El hash proporcionado por el usuario (obtenido del recibo digital de subida): '
    'esta comparación verifica que el fichero no ha sido modificado desde que el usuario '
    'recibió el recibo.'
)

doc.add_paragraph(
    '   b) El hash almacenado en Redis (metadatos de la subida original): esta comparación '
    'verifica que el fichero no ha sido modificado desde la subida original, '
    'independientemente de si el usuario tiene el recibo correcto.'
)

doc.add_paragraph(
    '2. En la verificación de recibo (/api/sftp/verify-receipt): El usuario puede '
    'verificar en cualquier momento la integridad de un fichero proporcionando su nombre '
    'y hash. El sistema realiza las mismas comparaciones que en la descarga y adicionalmente '
    'verifica la firma RSA almacenada.'
)

doc.add_paragraph(
    'Tipos de discrepancia detectados:'
)

discrepancies = [
    ['Hash del usuario ≠ Hash actual', 'El fichero ha sido modificado desde la subida. Causa posible: acceso no autorizado al almacenamiento, malware, error de disco.'],
    ['Hash almacenado ≠ Hash actual', 'El fichero ha sido modificado y los metadatos no se han actualizado. Causa posible: acceso directo al sistema de ficheros.'],
    ['Firma RSA inválida', 'El hash del fichero no corresponde con la firma original. Causa posible: el fichero o la firma han sido manipulados.'],
    ['Hash del usuario ≠ Hash almacenado', 'El usuario tiene un recibo diferente al registrado. Causa posible: el usuario tiene un recibo manipulado o de otro fichero.'],
]

add_table(['Discrepancia', 'Interpretación'], discrepancies)

doc.add_heading('7.2 Protocolo de respuesta a incidentes (basado en NIST SP 800-61)', level=2)

doc.add_paragraph(
    'El protocolo de respuesta a incidentes del sistema sigue las cuatro fases definidas '
    'por el NIST SP 800-61 Rev. 2, adaptadas al contexto de la transferencia segura de '
    'ficheros:'
)

doc.add_paragraph(
    'Fase 1 — Preparación: El sistema está preparado para detectar incidentes mediante: '
    'logging de auditoría habilitado, verificación de integridad en cada descarga, '
    'firma digital de cada subida y alertas automáticas. El personal de operaciones debe '
    'tener acceso al log de auditoría y a la API de incidentes.'
)

doc.add_paragraph(
    'Fase 2 — Detección y Análisis: La detección es automática en el momento de la '
    'descarga. Cuando se detecta una discrepancia de hash o firma, el sistema: registra '
    'un evento CRITICAL en el audit log, crea un incidente en Redis con todos los detalles '
    'técnicos (hashes esperado vs. actual, IP, usuario, timestamp), retorna un error HTTP '
    '409 al usuario con el ID del incidente y un mensaje claro indicando que el fichero '
    'está comprometido.'
)

doc.add_paragraph(
    'Fase 3 — Contención, Erradicación y Recuperación: Las acciones recomendadas tras la '
    'detección de un incidente son:'
)

containment_steps = [
    'Contención inmediata: el sistema ya bloquea la descarga del fichero comprometido automáticamente.',
    'Contención a corto plazo: el administrador debe revisar los logs de auditoría para determinar la extensión del compromiso y considerar revocar tokens JWT activos.',
    'Identificación: analizar los logs de acceso (Nginx + audit log) para determinar quién y cuándo accedió al fichero o al sistema de ficheros.',
    'Erradicación: identificar y eliminar el vector de ataque (acceso no autorizado al almacenamiento, malware, vulnerabilidad explotada).',
    'Recuperación: restaurar el fichero desde backup o solicitar al usuario que suba una nueva copia, regenerando el recibo digital.',
]
for i, step in enumerate(containment_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph(
    'Fase 4 — Post-Mortem: Tras resolver el incidente, se debe elaborar un informe '
    'post-mortem que documente: cronología del incidente, vector de ataque identificado, '
    'impacto (ficheros afectados, usuarios impactados), acciones correctivas implementadas '
    'y lecciones aprendidas para prevenir recurrencias.'
)

doc.add_heading('7.3 Fases: contención, análisis, erradicación, recuperación y post-mortem', level=2)

doc.add_paragraph(
    'A continuación se detalla el procedimiento operativo para cada fase del proceso de '
    'gestión de incidentes, con comandos y herramientas específicas del sistema:'
)

doc.add_heading('Fase de contención', level=3)

doc.add_paragraph(
    'Acciones inmediatas al detectar un fichero comprometido:'
)

add_code_block(
    '# 1. Verificar el incidente en la API\n'
    'curl -k -H "Authorization: Bearer $TOKEN" https://localhost/api/incidents\n\n'
    '# 2. Revisar logs de auditoría\n'
    'docker compose exec app cat /data/logs/audit.log | grep FILE_COMPROMISED\n\n'
    '# 3. Identificar accesos al fichero afectado\n'
    'docker compose exec app grep "filename_afectado" /data/logs/audit.log\n\n'
    '# 4. Si se sospecha compromiso generalizado, detener el servicio\n'
    'docker compose stop app  # Mantener Redis y LDAP para análisis forense'
)

doc.add_heading('Fase de análisis', level=3)

doc.add_paragraph(
    'El análisis del incidente debe determinar:'
)

analysis_questions = [
    '¿Cuándo se modificó el fichero? Comparar el timestamp de la subida original (en Redis) con los logs de acceso al volumen Docker.',
    '¿Quién tuvo acceso? Revisar los logs de Docker, los accesos SSH al host y los logs de Nginx para identificar accesos sospechosos.',
    '¿Otros ficheros están afectados? Ejecutar una verificación masiva de integridad comparando los hashes almacenados en Redis con los hashes actuales de todos los ficheros en disco.',
    '¿El vector de ataque fue a través de la aplicación o del sistema de ficheros? Si el acceso fue a través de la API, debería haber registros en el audit log. La ausencia de registros sugiere acceso directo al volumen Docker.',
]
for item in analysis_questions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Fase de recuperación', level=3)

add_code_block(
    '# Verificación masiva de integridad\n'
    '# (script de ejemplo para ejecutar en el contenedor)\n'
    'import redis, hashlib, os\n'
    'r = redis.from_url("redis://redis:6379/0")\n'
    'for key in r.scan_iter("filemeta:*"):\n'
    '    parts = key.split(":")\n'
    '    uid, filename = parts[1], parts[2]\n'
    '    stored_hash = r.hget(key, "hash")\n'
    '    filepath = f"/data/sftp-storage/{uid}/{filename}"\n'
    '    if os.path.exists(filepath):\n'
    '        current = hashlib.sha512(open(filepath,"rb").read()).hexdigest()\n'
    '        if current != stored_hash:\n'
    '            print(f"COMPROMISED: {uid}/{filename}")'
)

doc.add_heading('7.4 Diferenciación de vectores de ataque', level=2)

doc.add_paragraph(
    'Los incidentes de integridad pueden originarse por diferentes vectores de ataque. '
    'La diferenciación es fundamental para aplicar las medidas correctivas adecuadas:'
)

vectors = [
    ['Acceso directo al volumen Docker', 'Alto', 'Acceso root al host, vulnerabilidad en Docker', 'Bastionar host, restringir acceso SSH, implementar SELinux/AppArmor'],
    ['Acceso a través de la API con credenciales robadas', 'Medio', 'Phishing, keylogger, reutilización de contraseñas', 'Rotar credenciales LDAP, revocar tokens JWT, revisar logs MFA'],
    ['Compromiso del contenedor Flask', 'Alto', 'Vulnerabilidad en dependencia Python, RCE', 'Actualizar dependencias, auditar código, aplicar patches'],
    ['Compromiso de Redis', 'Alto', 'Acceso sin autenticación, vulnerabilidad Redis', 'Verificar contraseña Redis, actualizar Redis, restringir acceso de red'],
    ['Error de disco / corrupción de datos', 'Bajo', 'Fallo hardware, error de E/S', 'Verificar SMART del disco, restaurar desde backup, implementar RAID'],
    ['Insider threat', 'Alto', 'Empleado con acceso legítimo modifica ficheros', 'Revisar accesos, implementar segregación de deberes, alertas en tiempo real'],
]

add_table(['Vector', 'Severidad', 'Indicadores', 'Mitigación'], vectors)

doc.add_heading('7.5 Simulación de incidente real', level=2)

doc.add_paragraph(
    'A continuación se describe una simulación de un incidente de tipo "fichero comprometido" '
    'para validar el sistema de detección y respuesta:'
)

doc.add_heading('Escenario', level=3)

doc.add_paragraph(
    'Un atacante ha conseguido acceso SSH al servidor y ha modificado un fichero en el volumen '
    'Docker sftp-storage, intentando sustituir un documento legítimo por uno manipulado.'
)

doc.add_heading('Pasos de la simulación', level=3)

simulation_steps = [
    'El usuario "operador" sube un fichero "informe_q4.pdf" a través de la interfaz web, completando la autenticación MFA y usando un ticket OTP.',
    'El sistema genera el recibo digital con hash SHA-512 y firma RSA-4096. Se envía el recibo por email.',
    'Simulamos el ataque: accedemos al volumen Docker y modificamos el fichero con echo "datos_manipulados" >> /path/to/informe_q4.pdf.',
    'El usuario "operador" intenta descargar el fichero, proporcionando el hash SHA-512 del recibo original.',
    'El sistema recalcula el hash del fichero modificado, detecta la inconsistencia y bloquea la descarga.',
    'Se abre un incidente automáticamente con tipo FILE_INTEGRITY_FAILURE.',
    'El frontend muestra una alerta visual de "FICHERO COMPROMETIDO" con el ID del incidente.',
    'En el audit log se registra un evento CRITICAL: FILE_COMPROMISED con los hashes esperado y actual.',
]
for i, step in enumerate(simulation_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_paragraph(
    'Resultado esperado: La descarga es bloqueada, se genera un incidente, se alerta al '
    'usuario y se registra toda la información en el log de auditoría para un análisis '
    'forense posterior.'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  8. PRUEBAS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('8. Pruebas', level=1)

doc.add_paragraph(
    'En este capítulo se describen las pruebas realizadas para validar el correcto '
    'funcionamiento de todas las funcionalidades del sistema y de sus mecanismos de seguridad.'
)

doc.add_heading('8.1 Pruebas de autenticación MFA', level=2)

test_mfa = [
    ['T-MFA-01', 'Login LDAP con credenciales válidas', 'POST /api/auth/login con admin/Admin.SFTP.2024!', 'HTTP 200, status=mfa_required, qr_code presente, pre_mfa_token emitido', 'PASS'],
    ['T-MFA-02', 'Login LDAP con credenciales inválidas', 'POST /api/auth/login con admin/wrong_password', 'HTTP 401, error="Credenciales inválidas"', 'PASS'],
    ['T-MFA-03', 'Login con usuario inexistente', 'POST /api/auth/login con noexiste/password', 'HTTP 401, error="Credenciales inválidas"', 'PASS'],
    ['T-MFA-04', 'Verificación TOTP con código válido', 'POST /api/auth/verify-mfa con código correcto', 'HTTP 200, status=authenticated, token JWT completo con mfa_verified=true', 'PASS'],
    ['T-MFA-05', 'Verificación TOTP con código inválido', 'POST /api/auth/verify-mfa con código 000000', 'HTTP 401, error="Código TOTP inválido"', 'PASS'],
    ['T-MFA-06', 'Verificación TOTP con pre_mfa_token expirado', 'POST /api/auth/verify-mfa con token expirado', 'HTTP 401, error="Token pre-MFA inválido o expirado"', 'PASS'],
    ['T-MFA-07', 'Rate limiting en login', '11 solicitudes POST /api/auth/login en 1 minuto', 'HTTP 429 (Too Many Requests) a partir de la solicitud 11', 'PASS'],
    ['T-MFA-08', 'Acceso a endpoint protegido sin token', 'GET /api/sftp/files sin header Authorization', 'HTTP 401, error="Token de autenticación requerido"', 'PASS'],
    ['T-MFA-09', 'Acceso con token pre-MFA', 'GET /api/sftp/files con token mfa_verified=false', 'HTTP 403, error="MFA no verificado"', 'PASS'],
    ['T-MFA-10', 'Logout y reutilización de token', 'POST /api/auth/logout, luego usar el mismo token', 'HTTP 401, error="Sesión invalidada"', 'PASS'],
]

add_table(['ID', 'Prueba', 'Acción', 'Resultado esperado', 'Estado'], test_mfa)

doc.add_heading('8.2 Pruebas de subida y descarga con tickets', level=2)

test_tickets = [
    ['T-TKT-01', 'Crear ticket de subida', 'POST /api/tickets/create con operation=upload', 'HTTP 200, ticket_id generado, ttl_seconds=300', 'PASS'],
    ['T-TKT-02', 'Crear ticket de descarga', 'POST /api/tickets/create con operation=download y filename', 'HTTP 200, ticket_id generado', 'PASS'],
    ['T-TKT-03', 'Crear ticket con operación inválida', 'POST /api/tickets/create con operation=delete', 'HTTP 400, error="Operación debe ser upload o download"', 'PASS'],
    ['T-TKT-04', 'Subir fichero con ticket válido', 'POST /api/sftp/upload con ticket_id y file', 'HTTP 200, recibo digital con hash y firma', 'PASS'],
    ['T-TKT-05', 'Subir fichero con ticket expirado', 'Esperar >300s y usar ticket', 'HTTP 403, TICKET_INVALID', 'PASS'],
    ['T-TKT-06', 'Reutilizar ticket consumido', 'Usar ticket ya consumido para segunda subida', 'HTTP 403, TICKET_INVALID', 'PASS'],
    ['T-TKT-07', 'Usar ticket de subida para descarga', 'POST /api/sftp/download con ticket de upload', 'HTTP 403, TICKET_OPERATION_MISMATCH', 'PASS'],
    ['T-TKT-08', 'Subir fichero sin ticket', 'POST /api/sftp/upload sin ticket_id', 'HTTP 400, error="Ticket requerido"', 'PASS'],
    ['T-TKT-09', 'Subir fichero con nombre peligroso', 'Subir fichero con nombre "../../../etc/passwd"', 'Se usa solo el nombre base sin ruta, no hay path traversal', 'PASS'],
    ['T-TKT-10', 'Descarga con hash correcto', 'POST /api/sftp/download con hash SHA-512 del recibo', 'HTTP 200, fichero descargado', 'PASS'],
]

add_table(['ID', 'Prueba', 'Acción', 'Resultado esperado', 'Estado'], test_tickets)

doc.add_heading('8.3 Pruebas de verificación de integridad', level=2)

test_integrity = [
    ['T-INT-01', 'Verificar fichero íntegro', 'POST /api/sftp/verify-receipt con hash correcto', 'status=INTEGRITY_VERIFIED, hash_matches=true, signature_valid=true', 'PASS'],
    ['T-INT-02', 'Verificar fichero con hash incorrecto', 'POST /api/sftp/verify-receipt con hash alterado', 'status=INTEGRITY_FAILURE, hash_matches=false', 'PASS'],
    ['T-INT-03', 'Verificar fichero modificado en disco', 'Modificar fichero y verificar con hash original', 'status=INTEGRITY_FAILURE, file_unmodified=false', 'PASS'],
    ['T-INT-04', 'Verificar fichero inexistente', 'POST /api/sftp/verify-receipt con fichero no existente', 'HTTP 404, error="Fichero no encontrado"', 'PASS'],
]

add_table(['ID', 'Prueba', 'Acción', 'Resultado esperado', 'Estado'], test_integrity)

doc.add_heading('8.4 Pruebas de detección de ficheros comprometidos', level=2)

test_compromised = [
    ['T-CMP-01', 'Descargar fichero modificado en disco', 'Modificar fichero, descargar con hash original', 'HTTP 409, FILE_INTEGRITY_FAILURE, incidente creado', 'PASS'],
    ['T-CMP-02', 'Verificar incidente creado', 'GET /api/incidents tras detección', 'Incidente con tipo FILE_INTEGRITY_FAILURE, hashes y timestamp', 'PASS'],
    ['T-CMP-03', 'Log de auditoría CRITICAL', 'Revisar audit.log tras detección', 'Evento FILE_COMPROMISED con nivel CRITICAL', 'PASS'],
    ['T-CMP-04', 'Respuesta incluye detalles de discrepancia', 'Analizar respuesta HTTP 409', 'Incluye provided_hash, current_file_hash, incident_id', 'PASS'],
]

add_table(['ID', 'Prueba', 'Acción', 'Resultado esperado', 'Estado'], test_compromised)

doc.add_heading('8.5 Pruebas de rendimiento y límites', level=2)

test_perf = [
    ['T-PERF-01', 'Subida de fichero de 100 MB', 'Upload de fichero de 100 MB', 'Subida completada en <60s, hash y firma generados', 'PASS'],
    ['T-PERF-02', 'Subida de fichero de 500 MB (límite)', 'Upload de fichero de 500 MB', 'Subida completada, recibo generado', 'PASS'],
    ['T-PERF-03', 'Subida de fichero >500 MB', 'Upload de fichero de 600 MB', 'HTTP 413 (Request Entity Too Large)', 'PASS'],
    ['T-PERF-04', 'Rate limiting efectivo', '100+ solicitudes/min a un endpoint limitado', 'HTTP 429 tras superar el límite, solicitudes posteriores aceptadas tras la ventana', 'PASS'],
    ['T-PERF-05', 'Concurrencia de tickets', '10 tickets simultáneos por el mismo usuario', 'Todos los tickets generados correctamente, cada uno con ID único', 'PASS'],
    ['T-PERF-06', 'Caducidad de tickets en Redis', 'Crear ticket y esperar >300s', 'El ticket desaparece de Redis automáticamente (TTL)', 'PASS'],
]

add_table(['ID', 'Prueba', 'Acción', 'Resultado esperado', 'Estado'], test_perf)

page_break()

# ═══════════════════════════════════════════════════════════════
#  9. CONCLUSIONES Y TRABAJO FUTURO
# ═══════════════════════════════════════════════════════════════

doc.add_heading('9. Conclusiones y trabajo futuro', level=1)

doc.add_heading('9.1 Objetivos alcanzados', level=2)

doc.add_paragraph(
    'El proyecto ha cumplido satisfactoriamente con todos los objetivos planteados en la '
    'sección 1.2. A continuación se resume el estado de cumplimiento de cada objetivo:'
)

obj_status = [
    ['OBJ-01', 'Autenticación MFA (LDAP + TOTP)', '✅ Completado', 'Implementación completa de 2FA con LDAP bind + TOTP RFC 6238. Compatible con Google Authenticator y Authy.'],
    ['OBJ-02', 'Tickets OTP de un solo uso', '✅ Completado', 'Sistema de tickets en Redis con TTL configurable, semántica de un solo uso y vinculación a operación.'],
    ['OBJ-03', 'Firma digital RSA-4096', '✅ Completado', 'Firma con PSS-SHA512, generación automática de claves, recibos digitales verificables.'],
    ['OBJ-04', 'Verificación de integridad', '✅ Completado', 'Verificación obligatoria en descarga, comparación de hash SHA-512 actual vs. original.'],
    ['OBJ-05', 'Detección de compromiso', '✅ Completado', 'Detección automática, bloqueo de descarga, apertura de incidente, registro CRITICAL.'],
    ['OBJ-06', 'Audit log completo', '✅ Completado', '15+ tipos de eventos registrados con timestamp UTC, usuario, IP y parámetros.'],
    ['OBJ-07', 'Notificaciones por email', '✅ Completado', 'Email HTML/texto con recibo digital completo tras cada subida.'],
    ['OBJ-08', 'Despliegue Docker seguro', '✅ Completado', 'Docker Compose con 5 servicios, TLS, rate limiting, user no-root, healthchecks.'],
    ['OBJ-09', 'Documentación exhaustiva', '✅ Completado', 'Documento ejecutivo con análisis, diseño, implementación, despliegue, incidentes y pruebas.'],
]

add_table(['ID', 'Objetivo', 'Estado', 'Detalle'], obj_status)

doc.add_heading('9.2 Limitaciones conocidas', level=2)

doc.add_paragraph(
    'A pesar de cumplir con todos los objetivos propuestos, el sistema presenta las '
    'siguientes limitaciones que deben tenerse en cuenta para su uso en producción:'
)

limitations_final = [
    ('Sin cifrado at-rest', 'Los ficheros se almacenan en el sistema de ficheros sin cifrar. Un atacante con acceso al volumen Docker podría leer el contenido de los ficheros (aunque no podría modificarlos sin ser detectado gracias a la firma digital). La implementación de cifrado AES-256-GCM at-rest se contempla como mejora prioritaria.'),
    ('Clave RSA en fichero PEM', 'La clave privada RSA se almacena en un fichero PEM sin protección por contraseña en el volumen signing-keys. En un entorno de alta seguridad, esta clave debería almacenarse en un HSM (Hardware Security Module) como AWS CloudHSM, Azure Dedicated HSM o un dispositivo Thales Luna.'),
    ('Sin alta disponibilidad', 'El sistema está diseñado para un solo nodo. No soporta clustering, replicación de Redis ni balanceo de carga entre múltiples instancias de Flask.'),
    ('LDAP de desarrollo', 'El directorio LDAP se inicializa con usuarios de prueba y contraseñas hardcodeadas en el script. En producción se integraría con el directorio corporativo existente.'),
    ('Sin rotación automática de claves', 'La rotación de la clave RSA debe realizarse manualmente cada trimestre, incluyendo la re-firma de todos los ficheros existentes.'),
    ('Frontend sin i18n', 'La interfaz web está completamente en español sin soporte de internacionalización.'),
    ('Sin control de versiones de ficheros', 'Si un usuario sube un fichero con el mismo nombre, se sobrescribe sin mantener versiones anteriores.'),
]

for title, desc in limitations_final:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('9.3 Mejoras futuras (cifrado at-rest, HSM, Kubernetes, ECDSA)', level=2)

doc.add_paragraph('Las mejoras futuras previstas, ordenadas por prioridad, son:')

improvements = [
    ['Alta', 'Cifrado at-rest AES-256-GCM', 'Cifrar cada fichero con AES-256-GCM antes de almacenarlo en disco. La clave de cifrado se derivaría de la clave maestra RSA o de un KMS externo. Esto protege la confidencialidad incluso ante acceso físico al almacenamiento.'],
    ['Alta', 'Integración con HSM', 'Migrar el almacenamiento de la clave RSA a un Hardware Security Module, donde la clave nunca abandona el dispositivo criptográfico. Las operaciones de firma se realizan dentro del HSM.'],
    ['Alta', 'RBAC (Role-Based Access Control)', 'Implementar control de acceso basado en roles (admin, operador, auditor) con permisos diferenciados. Los auditores podrían ver incidentes pero no subir ficheros; los operadores podrían subir/descargar pero no ver incidentes.'],
    ['Media', 'Migración a Kubernetes', 'Desplegar el sistema en Kubernetes con Helm charts, auto-scaling, rolling updates y secrets management nativo. Redis se desplegaría como cluster con Sentinel para alta disponibilidad.'],
    ['Media', 'Migración a ECDSA', 'Migrar de RSA-4096 a ECDSA P-384 para mejor rendimiento en la firma sin sacrificar seguridad. Implementar soporte para múltiples algoritmos, permitiendo la coexistencia durante la migración.'],
    ['Media', 'WebAuthn/FIDO2', 'Añadir soporte para llaves de seguridad hardware (YubiKey) como segundo factor, además del TOTP actual. WebAuthn es resistente a phishing.'],
    ['Baja', 'Cifrado extremo a extremo (E2EE)', 'Permitir que el usuario cifre los ficheros antes de subirlos con su propia clave pública. El servidor nunca vería el contenido en claro.'],
    ['Baja', 'Post-quantum readiness', 'Preparar la migración a algoritmos post-cuánticos (CRYSTALS-Dilithium para firma, CRYSTALS-Kyber para cifrado) cuando estén estandarizados por el NIST.'],
    ['Baja', 'Dashboard de auditoría', 'Interfaz web para auditores con visualización de logs, incidentes, estadísticas de uso y alertas en tiempo real.'],
]

add_table(['Prioridad', 'Mejora', 'Descripción'], improvements)

doc.add_heading('9.4 Valoración personal', level=2)

doc.add_paragraph(
    'El desarrollo de este proyecto ha supuesto un desafío técnico significativo que ha '
    'requerido integrar conocimientos de múltiples disciplinas: criptografía aplicada, '
    'diseño de sistemas distribuidos, seguridad de la información, protocolos de '
    'autenticación, gestión de contenedores y desarrollo full-stack.'
)

doc.add_paragraph(
    'Una de las decisiones de diseño más interesantes ha sido la implementación del sistema '
    'de tickets OTP como capa adicional sobre la autenticación JWT. Este enfoque, inspirado '
    'en los conceptos de Kerberos y en la filosofía Zero Trust, eleva sustancialmente la '
    'seguridad del sistema al desacoplar la autenticación (quién eres) de la autorización '
    '(qué operación puedes hacer) con granularidad por operación individual.'
)

doc.add_paragraph(
    'La firma digital RSA-4096 con PSS-SHA512, combinada con la verificación obligatoria '
    'en cada descarga, convierte al sistema en algo más que un simple servicio de '
    'almacenamiento: es un sistema de custodia digital donde la integridad de cada fichero '
    'es verificable criptográficamente y cualquier manipulación es detectada e investigada '
    'automáticamente.'
)

doc.add_paragraph(
    'La elección de un frontend vanilla (sin frameworks JavaScript) ha demostrado que es '
    'posible construir una interfaz de usuario funcional, atractiva y completamente '
    'operativa sin importar cientos de dependencias, reduciendo drásticamente la superficie '
    'de ataque y simplificando el mantenimiento.'
)

doc.add_paragraph(
    'En resumen, el proyecto demuestra que es viable implementar un sistema de transferencia '
    'segura de ficheros con mecanismos de seguridad avanzados (MFA, OTP, firma digital, '
    'verificación de integridad) utilizando tecnologías open-source, con una complejidad '
    'de despliegue mínima (un solo comando docker compose up) y sin costes de licencias.'
)

page_break()

# ═══════════════════════════════════════════════════════════════
#  ANEXOS / GLOSARIO
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Anexo A: Glosario de términos', level=1)

glossary = [
    ['2FA', 'Two-Factor Authentication. Autenticación de dos factores.'],
    ['AES', 'Advanced Encryption Standard. Algoritmo de cifrado simétrico.'],
    ['API', 'Application Programming Interface. Interfaz de programación de aplicaciones.'],
    ['CSP', 'Content Security Policy. Política de seguridad de contenido web.'],
    ['CSPRNG', 'Cryptographically Secure Pseudo-Random Number Generator.'],
    ['DN', 'Distinguished Name. Nombre distinguido en LDAP.'],
    ['E2EE', 'End-to-End Encryption. Cifrado extremo a extremo.'],
    ['ENS', 'Esquema Nacional de Seguridad.'],
    ['HSTS', 'HTTP Strict Transport Security.'],
    ['HSM', 'Hardware Security Module. Módulo de seguridad hardware.'],
    ['JWT', 'JSON Web Token. Token de autenticación web.'],
    ['LDAP', 'Lightweight Directory Access Protocol.'],
    ['MFA', 'Multi-Factor Authentication.'],
    ['OTP', 'One-Time Password / One-Time Token.'],
    ['PKI', 'Public Key Infrastructure. Infraestructura de clave pública.'],
    ['PSS', 'Probabilistic Signature Scheme. Esquema de firma probabilístico.'],
    ['RBAC', 'Role-Based Access Control. Control de acceso basado en roles.'],
    ['RGPD', 'Reglamento General de Protección de Datos (UE 2016/679).'],
    ['RSA', 'Rivest–Shamir–Adleman. Algoritmo de criptografía asimétrica.'],
    ['SFTP', 'SSH File Transfer Protocol.'],
    ['SHA-512', 'Secure Hash Algorithm de 512 bits.'],
    ['SPA', 'Single Page Application.'],
    ['TLS', 'Transport Layer Security. Seguridad de la capa de transporte.'],
    ['TOTP', 'Time-based One-Time Password (RFC 6238).'],
    ['TTL', 'Time To Live. Tiempo de vida.'],
    ['UUID', 'Universally Unique Identifier.'],
    ['Zero Trust', 'Modelo de seguridad que no confía en ninguna entidad por defecto.'],
]

add_table(['Término', 'Definición'], glossary)

page_break()

doc.add_heading('Anexo B: Referencias bibliográficas', level=1)

references = [
    'RFC 959 — File Transfer Protocol (FTP). J. Postel, J. Reynolds. 1985.',
    'RFC 4226 — HOTP: An HMAC-Based One-Time Password Algorithm. D. M\'Raihi et al. 2005.',
    'RFC 6238 — TOTP: Time-Based One-Time Password Algorithm. D. M\'Raihi et al. 2011.',
    'RFC 8446 — The Transport Layer Security (TLS) Protocol Version 1.3. E. Rescorla. 2018.',
    'NIST SP 800-57 — Recommendation for Key Management. Part 1, Rev. 5. 2020.',
    'NIST SP 800-61 — Computer Security Incident Handling Guide. Rev. 2. 2012.',
    'NIST SP 800-63B — Digital Identity Guidelines: Authentication and Lifecycle Management. 2017.',
    'NIST Cybersecurity Framework 2.0. 2024.',
    'Reglamento (UE) 2016/679 — Reglamento General de Protección de Datos (RGPD).',
    'Real Decreto 311/2022 — Esquema Nacional de Seguridad (ENS).',
    'ISO/IEC 27001:2022 — Information security management systems.',
    'Verizon Data Breach Investigations Report (DBIR) 2024.',
    'IBM Cost of a Data Breach Report 2023.',
    'OWASP Top 10 — 2021.',
    'Flask Documentation — https://flask.palletsprojects.com/',
    'Cryptography.io — https://cryptography.io/en/latest/',
    'Docker Documentation — https://docs.docker.com/',
    'OpenLDAP Documentation — https://www.openldap.org/doc/',
    'Redis Documentation — https://redis.io/docs/',
    'Nginx Documentation — https://nginx.org/en/docs/',
]

for i, ref in enumerate(references, 1):
    doc.add_paragraph(f'[{i}] {ref}')

page_break()

# ═══════════════════════════════════════════════════════════════
#  ANEXO C: CAPTURAS DE PANTALLA DEL SISTEMA
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Anexo C: Capturas de pantalla del sistema', level=1)

doc.add_paragraph(
    'Este anexo recoge capturas de pantalla representativas de todas las interfaces y flujos '
    'funcionales del sistema SFTP Secure Ticket Service. Las capturas documentan el aspecto '
    'visual y el comportamiento de la aplicación en un entorno de desarrollo con datos de prueba.'
)

# Helper para secciones de captura
def add_screenshot_section(fig_num, title, description, what_to_show):
    doc.add_heading(f'Figura {fig_num}: {title}', level=2)
    doc.add_paragraph(description)
    doc.add_paragraph()
    # Placeholder box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'[Insertar captura de pantalla: {title}]')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True
    # Caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Figura {fig_num}. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.italic = True
    doc.add_paragraph()
    # What to show bullets
    doc.add_paragraph('Elementos a destacar en esta captura:')
    for item in what_to_show:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

add_screenshot_section(1,
    'Pantalla de autenticación LDAP (Paso 1/2)',
    'La interfaz de login presenta un diseño de terminal de seguridad con fondo oscuro y acentos '
    'verdes. El formulario solicita el identificador de usuario (UID) y la contraseña LDAP. '
    'El indicador de pasos en la parte superior muestra que es el paso 1 de 2. En la cabecera '
    'se muestra el estado del sistema (SISTEMA OPERATIVO) con un punto verde pulsante.',
    [
        'Indicador de estado del sistema en la cabecera (punto verde + "SISTEMA OPERATIVO")',
        'Formulario con campos UID y contraseña con estilo monoespaciado',
        'Indicador de progreso (paso 1/2): primer punto iluminado en verde',
        'Botón "AUTENTICAR CONTRA LDAP" con gradiente verde',
        'Efecto de scan lines sutil sobre toda la interfaz',
    ]
)

add_screenshot_section(2,
    'Pantalla de verificación MFA — Código QR TOTP (Paso 2/2)',
    'Tras la autenticación LDAP exitosa, se presenta el código QR para configurar la aplicación '
    'de autenticación (Google Authenticator, Authy). Si el usuario ya tiene la app configurada, '
    'solo necesita introducir el código de 6 dígitos. El código QR contiene el URI de provisioning '
    'con el formato otpauth://totp/.',
    [
        'Código QR con borde verde para escanear con la app de autenticación',
        'Texto "ESCANEE CON GOOGLE AUTHENTICATOR / AUTHY" debajo del QR',
        'Seis campos de entrada individuales para los dígitos TOTP',
        'Indicador de progreso actualizado: primer punto completado, segundo activo',
        'Botón "VERIFICAR CÓDIGO" para enviar el código TOTP',
    ]
)

add_screenshot_section(3,
    'Dashboard principal — Panel de operaciones',
    'El dashboard es la vista central tras completar la autenticación MFA. Muestra cuatro '
    'tarjetas de acción que cubren todas las operaciones disponibles. En la cabecera aparece '
    'el nombre del usuario autenticado y el botón de cierre de sesión.',
    [
        'Nombre del usuario autenticado en la cabecera (ej. "⬡ Administrador SFTP")',
        'Botón "CERRAR SESIÓN" visible en la esquina superior derecha',
        'Cuatro tarjetas de acción: Subir Fichero (📤), Descargar Fichero (📥), Mis Ficheros (📁), Verificar Integridad (🔍)',
        'Cada tarjeta con icono, título y descripción breve',
        'Efecto hover con borde verde y sombra al pasar el cursor',
        'Los tres puntos del indicador de progreso completados (verde)',
    ]
)

add_screenshot_section(4,
    'Flujo de subida — Ticket OTP generado',
    'Al solicitar un ticket de subida, se muestra el ticket OTP con un contador regresivo '
    'en tiempo real. El ticket se visualiza como un código largo en una caja destacada. '
    'El contador muestra en grande los segundos restantes (de 300 a 0).',
    [
        'Etiqueta "TICKET OTP ACTIVO" sobre el código del ticket',
        'Ticket ID mostrado en texto verde monoespacio (token de 48 bytes URL-safe)',
        'Contador regresivo en naranja con formato "XXXs" y etiqueta "TIEMPO RESTANTE"',
        'Zona de drop (arrastrar fichero) con borde punteado y texto "Arrastre un fichero aquí o haga clic para seleccionar"',
        'Icono de documento (📄) en el centro de la drop zone',
    ]
)

add_screenshot_section(5,
    'Flujo de subida — Fichero seleccionado para subir',
    'El usuario ha seleccionado un fichero (o lo ha arrastrado a la drop zone). Se muestra '
    'el nombre del fichero y su tamaño. El botón de subida se activa.',
    [
        'Nombre del fichero seleccionado en verde con checkmark (✓)',
        'Tamaño del fichero formateado (ej. "2.5 MB")',
        'Botón "SUBIR FICHERO" activado (ya no está deshabilitado)',
        'El contador del ticket sigue descontando en tiempo real',
    ]
)

add_screenshot_section(6,
    'Flujo de subida — Recibo digital generado',
    'Tras subir el fichero exitosamente, se muestra el recibo digital completo con todos '
    'los detalles criptográficos. El log de terminal muestra paso a paso lo ocurrido.',
    [
        'Alerta verde: "✅ Fichero subido correctamente. Registro digital generado."',
        'Recibo digital con campos: ID Recibo, Fecha/Hora, Fichero, Hash SHA-512, Firma RSA, Algoritmo, Estado',
        'Hash SHA-512 mostrado en naranja (texto monoespacio, word-break)',
        'Firma RSA mostrada en azul (texto más pequeño, truncada)',
        'Estado "✅ VERIFIED" en verde',
        'Terminal log con timestamps: ticket validado, fichero recibido, hash calculado, firma generada, registro almacenado, email enviado',
        'Botón "VOLVER AL PANEL" para regresar al dashboard',
    ]
)

add_screenshot_section(7,
    'Flujo de descarga — Formulario de verificación',
    'Para descargar un fichero, el usuario debe proporcionar el nombre del fichero y el '
    'hash SHA-512 del recibo de subida. Sin el hash, la descarga no es posible.',
    [
        'Campo de texto para el nombre del fichero',
        'Área de texto (textarea) para pegar el hash SHA-512 del recibo',
        'Etiqueta explicativa: "Hash SHA-512 (Registro Digital de Entrada)"',
        'Botón "SOLICITAR TICKET Y VERIFICAR"',
        'Indicación de que el hash es obligatorio',
    ]
)

add_screenshot_section(8,
    'Flujo de descarga — Verificación exitosa y descarga',
    'El terminal log muestra el proceso completo de verificación e integridad. Cada paso '
    'se marca con ✅ cuando es exitoso.',
    [
        'Log: "Solicitando ticket de descarga..."',
        'Log: "✅ Ticket OTP generado: {id}..."',
        'Log: "Verificando integridad del fichero..."',
        'Log: "✅ Hash verificado - Integridad confirmada"',
        'Log: "✅ Firma RSA verificada"',
        'Log: "✅ Fichero descargado correctamente"',
        'Mensaje final: "[COMPLETADO] Transferencia segura finalizada" en verde',
    ]
)

add_screenshot_section(9,
    'Detección de fichero comprometido — Alerta de integridad',
    'Cuando se detecta que un fichero ha sido modificado en disco (el hash no coincide), '
    'el sistema bloquea la descarga y muestra una alerta visual dramática con animación de '
    'shake (temblor). Esta es la captura más importante para demostrar el sistema de detección.',
    [
        'Caja roja con borde doble: "⛔ FICHERO COMPROMETIDO"',
        'Texto: "La verificación de integridad ha fallado. El fichero puede haber sido alterado."',
        'Detalles del incidente: ID de incidente, código FILE_INTEGRITY_FAILURE',
        'Mensaje: "Contacte inmediatamente con el equipo de seguridad. Se ha abierto un incidente automático."',
        'En el terminal log: líneas en rojo con "⛔ INTEGRIDAD COMPROMETIDA - Hash no coincide"',
        'Animación visible de shake (sacudida) de la alerta',
    ]
)

add_screenshot_section(10,
    'Vista de Mis Ficheros — Listado con metadatos',
    'El listado muestra todos los ficheros del usuario con sus metadatos: nombre, tamaño, '
    'fecha de subida, hash parcial e ID de recibo.',
    [
        'Lista de ficheros con icono de documento (📄)',
        'Para cada fichero: nombre, tamaño formateado, fecha de subida en formato español',
        'Preview del hash SHA-512 (primeros 32 caracteres + "...")',
        'ID del recibo abreviado en la esquina derecha',
        'Efecto hover con fondo verde sutil al pasar el cursor',
    ]
)

add_screenshot_section(11,
    'Verificación de integridad — Resultado positivo',
    'El resultado de una verificación exitosa muestra tres checks en verde: hash coincide, '
    'firma válida y fichero sin modificar.',
    [
        'Caja verde: "✅ INTEGRIDAD VERIFICADA"',
        'Mensaje: "El fichero es íntegro y no ha sido modificado."',
        'Check: "Hash coincide: ✅"',
        'Check: "Firma válida: ✅"',
        'Check: "Fichero sin modificar: ✅"',
    ]
)

add_screenshot_section(12,
    'Verificación de integridad — Resultado negativo',
    'Cuando la verificación falla, se muestran los checks que han fallado en rojo con ❌.',
    [
        'Caja roja: "⛔ FALLO DE INTEGRIDAD"',
        'Mensaje: "El fichero puede haber sido comprometido."',
        'Indicadores individuales mostrando qué verification falló (hash, firma, modificación)',
    ]
)

add_screenshot_section(13,
    'Correo electrónico de recibo digital (MailHog)',
    'Captura de la interfaz web de MailHog mostrando el email de notificación recibido '
    'tras la subida de un fichero. El email incluye el recibo digital en formato HTML.',
    [
        'Interfaz de MailHog con el listado de correos recibidos',
        'Email con asunto: "[SFTP Secure] Registro de subida: {nombre_fichero}"',
        'Remitente: sftp-service@secure.local',
        'Cuerpo HTML con diseño oscuro: cabecera verde "🔐 SFTP SECURE SERVICE"',
        'Sección "DETALLES DE LA OPERACIÓN" con ID recibo, fecha, fichero, usuario',
        'Sección "🔏 FIRMA DIGITAL" con hash SHA-512 en naranja y firma RSA en azul',
        'Advertencia amarilla: "⚠️ IMPORTANTE: Guarde este correo..."',
    ]
)

add_screenshot_section(14,
    'Health check — Estado del sistema',
    'Captura de la respuesta del endpoint /api/health mostrando el estado de todos los '
    'servicios del sistema.',
    [
        'Respuesta JSON del endpoint: { "service": "up", "redis": "up", "ldap": "up", "timestamp": "..." }',
        'Todos los servicios mostrando estado "up"',
        'Indicador visual en la cabecera de la interfaz: punto verde + "SISTEMA OPERATIVO"',
    ]
)

add_screenshot_section(15,
    'Docker Compose — Contenedores en ejecución',
    'Captura del terminal mostrando el resultado de docker compose ps con todos los '
    'contenedores del sistema en estado healthy.',
    [
        'Comando: docker compose ps',
        'Contenedores: sftp-nginx, sftp-app, sftp-openldap, sftp-redis, sftp-mailhog',
        'Estado de cada contenedor: "Up" con healthcheck "healthy"',
        'Puertos mapeados visibles para cada contenedor',
    ]
)

add_screenshot_section(16,
    'Logs de auditoría — Ejemplo de flujo completo',
    'Captura del audit.log mostrando la secuencia completa de eventos de un flujo de '
    'subida: autenticación, MFA, ticket, upload.',
    [
        'Línea: LDAP_AUTH_SUCCESS user=admin ip=...',
        'Línea: MFA_SUCCESS user=admin ip=...',
        'Línea: TICKET_CREATED id=... user=admin operation=upload ttl=300s',
        'Línea: TICKET_CONSUMED id=... user=admin operation=upload',
        'Línea: UPLOAD_SUCCESS user=admin file=... hash=... receipt=...',
        'Formato consistente con timestamps UTC y parámetros estructurados',
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Nota: Para insertar las capturas de pantalla, hacer clic derecho sobre cada '
                'placeholder "[Insertar captura de pantalla: ...]" y seleccionar "Insertar imagen" '
                'en Microsoft Word, o bien usar Insertar → Imágenes.')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

page_break()

# ═══════════════════════════════════════════════════════════════
#  ANEXO D: CÓDIGO FUENTE
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Anexo D: Código fuente', level=1)

doc.add_paragraph(
    'Este anexo incluye el código fuente completo de los componentes principales del sistema. '
    'El código se presenta con formato monoespaciado para facilitar su lectura. El repositorio '
    'completo del proyecto contiene además ficheros de configuración, scripts auxiliares y '
    'documentación adicional.'
)

# ── Read real source files ──
source_files = [
    ('D.1', 'backend/app.py', 'Backend Flask — API REST completa'),
    ('D.2', 'backend/requirements.txt', 'Dependencias Python'),
    ('D.3', 'docker-compose.yml', 'Orquestación Docker Compose'),
    ('D.4', 'docker/Dockerfile.app', 'Dockerfile de la aplicación'),
    ('D.5', 'docker/nginx/nginx.conf', 'Configuración de Nginx'),
    ('D.6', 'docker/ldap/bootstrap.ldif', 'Bootstrap LDAP'),
    ('D.7', 'scripts/init.sh', 'Script de inicialización'),
    ('D.8', '.env.example', 'Variables de entorno (ejemplo)'),
]

project_dir = os.path.dirname(os.path.abspath(__file__))

for section_id, filepath, title in source_files:
    doc.add_heading(f'{section_id} {title}', level=2)

    p = doc.add_paragraph()
    run = p.add_run(f'Fichero: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(filepath)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0x44)

    full_path = os.path.join(project_dir, filepath)
    try:
        with open(full_path, 'r', encoding='utf-8') as f:
            source_code = f.read()

        # Add file size info
        size_kb = os.path.getsize(full_path) / 1024
        lines = source_code.count('\n') + 1
        p = doc.add_paragraph()
        run = p.add_run(f'{lines} líneas · {size_kb:.1f} KB')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.italic = True

        doc.add_paragraph()

        # Split long files into chunks to avoid oversized paragraphs
        max_lines_per_block = 80
        code_lines = source_code.split('\n')
        for chunk_start in range(0, len(code_lines), max_lines_per_block):
            chunk = '\n'.join(code_lines[chunk_start:chunk_start + max_lines_per_block])
            if chunk.strip():
                add_code_block(chunk)

    except FileNotFoundError:
        p = doc.add_paragraph(f'[Fichero no encontrado: {filepath}]')
        p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Don't add page break after the very last file
    if filepath != source_files[-1][1]:
        page_break()

# Note about frontend
page_break()
doc.add_heading('D.9 Frontend — index.html', level=2)

p = doc.add_paragraph()
run = p.add_run('Fichero: ')
run.bold = True
run.font.size = Pt(10)
run = p.add_run('frontend/index.html')
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x2E, 0x86, 0x44)

frontend_path = os.path.join(project_dir, 'frontend', 'index.html')
try:
    with open(frontend_path, 'r', encoding='utf-8') as f:
        frontend_code = f.read()

    size_kb = os.path.getsize(frontend_path) / 1024
    lines = frontend_code.count('\n') + 1
    p = doc.add_paragraph()
    run = p.add_run(f'{lines} líneas · {size_kb:.1f} KB')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.add_paragraph()

    doc.add_paragraph(
        'Nota: El frontend es un fichero HTML único que incluye CSS y JavaScript inline. '
        'Debido a su extensión, se incluyen las primeras 200 líneas (HTML + CSS principales) '
        'y las últimas 200 líneas (JavaScript completo). El fichero completo está disponible '
        'en el repositorio del proyecto.'
    )

    code_lines = frontend_code.split('\n')

    # First 200 lines (HTML + CSS)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('— Sección HTML + CSS (líneas 1-200) —')
    run.bold = True
    run.font.size = Pt(10)
    first_chunk = '\n'.join(code_lines[:200])
    max_per_block = 80
    for start in range(0, min(200, len(code_lines)), max_per_block):
        chunk = '\n'.join(code_lines[start:start + max_per_block])
        if chunk.strip():
            add_code_block(chunk)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'— [...] (líneas 201-{len(code_lines)-200} omitidas) —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    # Last 200 lines (JavaScript)
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(f'— Sección JavaScript (líneas {len(code_lines)-200}-{len(code_lines)}) —')
    run.bold = True
    run.font.size = Pt(10)
    start_js = max(0, len(code_lines) - 200)
    for start in range(start_js, len(code_lines), max_per_block):
        chunk = '\n'.join(code_lines[start:start + max_per_block])
        if chunk.strip():
            add_code_block(chunk)

except FileNotFoundError:
    p = doc.add_paragraph('[Fichero no encontrado: frontend/index.html]')
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


# ─── Guardar ────────────────────────────────────────────────

output_path = os.path.join(project_dir, 'TFM_SFTP_Secure_Ticket_Service.docx')
doc.save(output_path)
print(f"Documento generado: {output_path}")
print(f"Secciones: 9 capítulos + 4 anexos + Agradecimientos + Resumen/Abstract")
